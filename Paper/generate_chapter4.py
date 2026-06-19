# -*- coding: utf-8 -*-
"""生成第4章 数据库逻辑结构设计 DOCX — 精修版
包含: ER转关系模式 / 范式规范化 / 视图 / 索引 / 存储过程&函数 / 触发器
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, sys

sys.stdout.reconfigure(encoding='utf-8', errors='replace')

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54); section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17); section.right_margin = Cm(3.17)

style = doc.styles['Normal']
style.font.name = '宋体'; style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Cm(0.74)

for lvl, sz in [(1, 16), (2, 14), (3, 13)]:
    s = doc.styles[f'Heading {lvl}']
    s.font.name = '黑体'; s.font.size = Pt(sz); s.font.bold = True
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if lvl == 1 else WD_ALIGN_PARAGRAPH.LEFT
    s.paragraph_format.space_before = Pt(12 if lvl == 1 else 8)
    s.paragraph_format.space_after = Pt(6 if lvl == 1 else 4)


def body(text):
    p = doc.add_paragraph(style='Normal'); p.clear()
    r = p.add_run(text)
    r.font.name = '宋体'; r.font.size = Pt(12)
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    return p


def code_block(sql_text):
    """Add a code block styled paragraph."""
    p = doc.add_paragraph(style='Normal'); p.clear()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(sql_text)
    r.font.name = 'Consolas'; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(30, 30, 30)
    return p


def make_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.font.name = '宋体'; r.font.size = Pt(10); r.font.bold = True
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), '4472C4'); shd.set(qn('w:val'), 'clear')
        c._tc.get_or_add_tcPr().append(shd)
        r.font.color.rgb = RGBColor(255, 255, 255)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.rows[i + 1].cells[j]; c.text = ''
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j < len(row) - 1 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val)); r.font.name = '宋体'; r.font.size = Pt(9)
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if i % 2 == 0:
                shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'F2F6FC'); shd.set(qn('w:val'), 'clear')
                c._tc.get_or_add_tcPr().append(shd)
    doc.add_paragraph()
    return t


def make_wide_table(headers, rows):
    """Table with left-aligned text for wide content."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.font.name = '宋体'; r.font.size = Pt(9); r.font.bold = True
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), '4472C4'); shd.set(qn('w:val'), 'clear')
        c._tc.get_or_add_tcPr().append(shd)
        r.font.color.rgb = RGBColor(255, 255, 255)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.rows[i + 1].cells[j]; c.text = ''
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val)); r.font.name = '宋体'; r.font.size = Pt(8.5)
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if i % 2 == 0:
                shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'F2F6FC'); shd.set(qn('w:val'), 'clear')
                c._tc.get_or_add_tcPr().append(shd)
    doc.add_paragraph()
    return t


def add_figure(img_path, caption, width=5.5):
    if not os.path.exists(img_path): return
    pi = doc.add_paragraph(); pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pi.paragraph_format.first_line_indent = Cm(0); pi.paragraph_format.space_before = Pt(6)
    pi.add_run().add_picture(img_path, width=Inches(width))
    pc = doc.add_paragraph(); pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pc.paragraph_format.first_line_indent = Cm(0)
    pc.paragraph_format.space_before = Pt(2); pc.paragraph_format.space_after = Pt(10)
    r = pc.add_run(caption); r.font.name = '宋体'; r.font.size = Pt(10)
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


# ================================================================
#  MATPLOTLIB SETUP (for any figures)
# ================================================================
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
plt.rcParams['font.sans-serif'] = ['SimHei']
plt.rcParams['font.family'] = 'sans-serif'
plt.rcParams['axes.unicode_minus'] = False

def save_fig(fig, path):
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close(fig)


# ================================================================
#  CHAPTER 4
# ================================================================
doc.add_heading('第4章  数据库逻辑结构设计', level=1)

body(
    '逻辑结构设计是将概念结构设计阶段产生的全局E-R图转换为特定数据库管理系统支持的'
    '关系数据模型的过程。本章的核心任务包括：遵循E-R图向关系模型的转换规则，将全部实体'
    '和联系映射为关系模式集合；对每个关系模式进行函数依赖分析和范式判定，通过规范化'
    '分解消除更新异常，确保所有关系模式达到第三范式（3NF）；在此基础上，设计数据视图'
    '以支持高频查询场景；建立索引以优化数据库检索性能；编写存储过程与自定义函数以封装'
    '核心业务逻辑；设计触发器以保障数据的完整性与一致性。'
)

# ===================== 4.1 E-R图向关系模型转换 =====================
doc.add_heading('4.1  E-R图向关系模型转换', level=2)

body(
    'E-R图向关系模型的转换遵循以下五条基本规则，以下结合全局E-R图中的实体和联系逐一'
    '应用这些规则，给出转换后的完整关系模式集合。在关系模式表示中，主键以下划线____标识，'
    '外键以波浪下划线标识。'
)

doc.add_heading('4.1.1  转换规则', level=3)

body(
    '规则一（实体转换）：每个实体类型转换为一个关系模式，实体的属性即为关系模式的属性，'
    '实体的标识符即为关系模式的主键。'
)

body(
    '规则二（1:1联系转换）：可以将1:1联系与任意一端的关系模式合并，在该端加入另一端'
    '关系模式的主键作为外键，同时将联系自身的属性（如有）一并加入。也可以将联系单独'
    '转换为一个独立的关系模式。本系统中，合同（Contract）与科研项目（Project）之间的'
    '1:1"签署"联系采用与Project端合并的方式。验收记录（Acceptance）与科研项目之间'
    '的1:1联系同理。'
)

body(
    '规则三（1:n联系转换）：将1端关系模式的主键作为外键加入n端关系模式中，同时将联系'
    '自身的属性（如有）也加入n端。本系统中包括：学院→研究人员（归属）、研究人员→项目'
    '（负责）、通知→项目（引导申报）、项目→预算（编制）、预算→经费流水（记录）、'
    '项目→成果（产出）、项目→进展报告（提交）、研究人员→通知（发布）等。'
)

body(
    '规则四（m:n联系转换）：m:n联系必须转换为一个独立的关系模式，其属性包括相联两端'
    '实体各自的主键（共同作为该关系模式的主键），以及联系自身的属性。本系统中，研究人员'
    '与科研项目之间的"评审"联系即为典型的m:n联系，转换为独立的评审记录（Review）关系模式。'
    '研究人员与科研项目之间的"参与"联系若单独建模，亦同此规则。本系统设计中，"参与"'
    '联系的信息已隐含在项目负责人字段及评审记录中，故未单设独立关系模式。'
)

body(
    '规则五（多元联系转换）：多元联系（三元及以上）同样转换为独立关系模式，其属性包括'
    '参与该联系的各实体主键及联系自身属性。本系统中未涉及三元及以上联系。'
)

doc.add_heading('4.1.2  转换后的关系模式集合', level=3)

body(
    '综合应用上述五条转换规则，将全局E-R图中的11类实体与14个联系映射为以下12个关系模式。'
    '其中，11个实体直接转换为11个关系模式，m:n"评审"联系转换为1个独立关系模式（Review），'
    '其余1:1和1:n联系均通过添加外键的方式合并到相应端的关系模式中。各关系模式详见表4-1。'
)

make_wide_table(
    ['序号', '关系模式名', '属性列表', '主键', '外键', '来源'],
    [
        ['R1', 'college',
         'college_id, name, dean, tel',
         'college_id', '—', '学院实体'],
        ['R2', 'researcher',
         'researcher_id, name, password, title, college_id, phone, email, role',
         'researcher_id', 'college_id→college', '研究人员实体 + 归属联系(1:n)'],
        ['R3', 'project',
         'project_id, name, type, level, leader_id, apply_date, start_date, end_date, budget_total, status, acceptance_date, file_url',
         'project_id', 'leader_id→researcher', '科研项目实体 + 负责联系(1:n)'],
        ['R4', 'budget',
         'budget_id, project_id, category, amount, spent',
         'budget_id', 'project_id→project', '预算科目实体 + 编制联系(1:n)'],
        ['R5', 'expenditure',
         'exp_id, project_id, budget_id, type, amount, exp_date, purpose, operator_id, approval_status',
         'exp_id', 'project_id→project, budget_id→budget, operator_id→researcher', '经费流水实体 + 记录/登记/报销联系(1:n)'],
        ['R6', 'review',
         'project_id, expert_id, result, score, comment, review_date',
         '(project_id, expert_id)', 'project_id→project, expert_id→researcher', '评审联系(m:n)独立转换'],
        ['R7', 'notice',
         'notice_id, title, content, publish_date, publisher_id',
         'notice_id', 'publisher_id→researcher', '通知公告实体 + 发布联系(1:n)'],
        ['R8', 'contract',
         'contract_id, project_id, file_url, sign_date, content',
         'contract_id', 'project_id→project', '合同实体 + 签署联系(1:1→合并到Project端，此处单列)'],
        ['R9', 'progress_report',
         'report_id, project_id, report_year, submit_date, content, review_status',
         'report_id', 'project_id→project', '进展报告实体 + 提交联系(1:n)'],
        ['R10', 'change_request',
         'change_id, project_id, change_type, old_value, new_value, reason, request_date, approval_status',
         'change_id', 'project_id→project', '变更申请实体(需求分析中识别)'],
        ['R11', 'acceptance',
         'acceptance_id, project_id, apply_date, review_result, acceptance_date, certificate_url',
         'acceptance_id', 'project_id→project', '验收记录实体 + 验收联系(1:1)'],
        ['R12', 'achievement',
         'ach_id, project_id, type, title, publish_date, author, file_url, review_status',
         'ach_id', 'project_id→project', '科研成果实体 + 产出联系(1:n)'],
    ]
)


# ===================== 4.2 关系模式规范化 =====================
doc.add_heading('4.2  关系模式规范化（3NF）', level=2)

body(
    '规范化是消除关系模式中数据冗余和更新异常的系统性方法。本节选取系统中具有代表性的'
    '关系模式，逐一进行函数依赖分析、范式判定和规范化分解（如需要），确保所有关系模式'
    '均达到第三范式（3NF）及以上。以下对四个核心关系模式进行详尽的规范化论证。'
)

# 4.2.1 Researcher
doc.add_heading('4.2.1  研究人员关系模式（researcher）', level=3)

body(
    '关系模式：R_researcher(researcher_id, name, password, title, college_id, phone, email, role)'
)

body(
    '函数依赖集 F = {'
    'researcher_id → name, '
    'researcher_id → password, '
    'researcher_id → title, '
    'researcher_id → college_id, '
    'researcher_id → phone, '
    'researcher_id → email, '
    'researcher_id → role'
    '}。所有非主属性均完全函数依赖于候选键 researcher_id，不存在部分函数依赖。'
)

body(
    '传递依赖分析：college_id → {college_name, dean, tel} 属于 college 关系模式内部的'
    '函数依赖，在 researcher 关系模式中 college_id 仅为外键引用，不存在 researcher_id → '
    'college_id → college_name 的传递依赖（因为 college_name 不在本关系模式中）。'
    '因此该关系模式不存在传递函数依赖。'
)

body(
    '结论：R_researcher 中每一个非主属性都完全函数依赖于候选键 researcher_id，且不存在'
    '传递函数依赖，满足第三范式（3NF）的定义。同时，由于候选键为单属性，不存在非主属性'
    '对候选键的部分函数依赖，自动满足第二范式（2NF）。该关系模式无需进一步分解。'
)

# 4.2.2 Project
doc.add_heading('4.2.2  科研项目关系模式（project）', level=3)

body(
    '关系模式：R_project(project_id, name, type, level, leader_id, apply_date, '
    'start_date, end_date, budget_total, status, acceptance_date, file_url)'
)

body(
    '函数依赖集 F = {'
    'project_id → name, '
    'project_id → type, '
    'project_id → level, '
    'project_id → leader_id, '
    'project_id → apply_date, '
    'project_id → start_date, '
    'project_id → end_date, '
    'project_id → budget_total, '
    'project_id → status, '
    'project_id → acceptance_date, '
    'project_id → file_url'
    '}。所有非主属性均完全函数依赖于单一候选键 project_id，不存在部分函数依赖。'
)

body(
    '传递依赖分析：虽然 leader_id → {leader_name, leader_title, ...} 存在于 researcher '
    '关系模式中，但在 project 关系模式内 leader_id 是外键，不构成 project_id → leader_id → '
    'leader_name 的传递依赖问题，因为 leader_name 等属性存储于 researcher 表而非 project 表中。'
    'status 字段虽然决定了项目的生命周期阶段，但 status 本身是非主属性，不构成对其他非主属性'
    '的函数决定关系。'
)

body(
    '结论：R_project 满足第三范式（3NF）要求。所有非主属性均直接依赖于主键 project_id，'
    '不存在部分依赖和传递依赖。该关系模式无需进一步分解。'
)

# 4.2.3 Expenditure
doc.add_heading('4.2.3  经费流水关系模式（expenditure）', level=3)

body(
    '关系模式：R_expenditure(exp_id, project_id, budget_id, type, amount, exp_date, '
    'purpose, operator_id, approval_status)'
)

body(
    '函数依赖集 F = {'
    'exp_id → project_id, '
    'exp_id → budget_id, '
    'exp_id → type, '
    'exp_id → amount, '
    'exp_id → exp_date, '
    'exp_id → purpose, '
    'exp_id → operator_id, '
    'exp_id → approval_status'
    '}。候选键为 exp_id（单属性），不存在部分函数依赖。'
)

body(
    '传递依赖分析：project_id、budget_id、operator_id 均为外键，它们各自决定的外部属性'
    '（如项目名称、预算科目名称、操作员姓名等）不存在于本关系模式中，因此不构成传递依赖。'
    'approval_status 虽为状态字段，但不函数决定其他非主属性。'
)

body(
    '结论：R_expenditure 满足第三范式（3NF）。该关系模式通过外键引用维护了与 project、'
    'budget 和 researcher 的关联，自身结构规范，无冗余和更新异常。'
)

# 4.2.4 Review
doc.add_heading('4.2.4  评审记录关系模式（review）', level=3)

body(
    '关系模式：R_review(project_id, expert_id, result, score, comment, review_date)'
)

body(
    '函数依赖集 F = {'
    '(project_id, expert_id) → result, '
    '(project_id, expert_id) → score, '
    '(project_id, expert_id) → comment, '
    '(project_id, expert_id) → review_date'
    '}。候选键为复合主键 (project_id, expert_id)，所有非主属性均完全函数依赖于该复合键，'
    '不存在部分函数依赖（因为没有任何非主属性仅依赖于 project_id 或 expert_id 中的一个）。'
)

body(
    '传递依赖分析：不存在某个非主属性函数决定另一个非主属性的情况。result字段取值为'
    '"通过"/"不通过"，score字段取值为0-100，两者相互独立，不存在 result → score 或 '
    'score → result 的函数依赖关系。'
)

body(
    '结论：R_review 满足第三范式（3NF）和BCNF（Boyce-Codd范式）。该关系模式的每一个'
    '决定因素都包含候选键，满足BCNF的严格定义，无需进一步分解。该模式是由m:n联系转换而来'
    '的典型关联关系模式，结构简洁且符合规范化要求。'
)

# ---- 范式规范化总结表 ----
doc.add_heading('4.2.5  规范化总结', level=3)

body(
    '经上述逐项分析，系统中的全部12个关系模式均至少满足第三范式（3NF）。表4-2给出'
    '所有关系模式的范式判定结果汇总。'
)

make_table(
    ['序号', '关系模式名', '候选键', '最高范式', '判定依据'],
    [
        ['R1', 'college', 'college_id', '3NF/BCNF', '单属性主键，无部分依赖，无传递依赖'],
        ['R2', 'researcher', 'researcher_id', '3NF', '单属性主键，外键不引入传递依赖'],
        ['R3', 'project', 'project_id', '3NF', '单属性主键，外键不引入传递依赖'],
        ['R4', 'budget', 'budget_id', '3NF', '单属性主键，project_id为外键'],
        ['R5', 'expenditure', 'exp_id', '3NF', '单属性主键，所有外键不引入传递依赖'],
        ['R6', 'review', '(project_id, expert_id)', 'BCNF', '复合主键，决定因素均为超键'],
        ['R7', 'notice', 'notice_id', '3NF', '单属性主键，publisher_id为外键'],
        ['R8', 'contract', 'contract_id', '3NF', '单属性主键，project_id为外键'],
        ['R9', 'progress_report', 'report_id', '3NF', '单属性主键，project_id为外键'],
        ['R10', 'change_request', 'change_id', '3NF', '单属性主键，project_id为外键'],
        ['R11', 'acceptance', 'acceptance_id', '3NF', '单属性主键，project_id为外键'],
        ['R12', 'achievement', 'ach_id', '3NF', '单属性主键，project_id为外键'],
    ]
)


# ===================== 4.3 视图设计 =====================
doc.add_heading('4.3  视图设计', level=2)

body(
    '数据视图（View）是基于基本表的虚拟表，用于简化复杂查询、提高数据安全性和实现逻辑'
    '数据独立性。本系统设计了两个核心视图，分别服务于项目预算执行监控和研究人员工作'
    '量统计两个高频查询场景。'
)

doc.add_heading('4.3.1  视图一：项目预算执行状态视图', level=3)

body(
    '视图名称：v_project_budget_status。功能说明：汇总每个项目的总预算金额、总支出金额、'
    '预算执行率（已支出/预算总额×100%）和剩余预算，为财务管理和科研处决策提供一站式查询。'
    '该视图通过JOIN project表和expenditure表实现数据的聚合展示。'
)

code_block(
    'CREATE VIEW v_project_budget_status AS\n'
    'SELECT\n'
    '    p.project_id,\n'
    '    p.name AS project_name,\n'
    '    p.budget_total,\n'
    '    COALESCE(SUM(e.amount), 0) AS total_spent,\n'
    '    ROUND(COALESCE(SUM(e.amount), 0) / p.budget_total * 100, 2) AS execution_rate,\n'
    '    p.budget_total - COALESCE(SUM(e.amount), 0) AS remaining_budget\n'
    'FROM project p\n'
    'LEFT JOIN expenditure e ON p.project_id = e.project_id\n'
    '    AND e.type = \'支出\' AND e.approval_status = \'已通过\'\n'
    'GROUP BY p.project_id, p.name, p.budget_total;'
)

doc.add_heading('4.3.2  视图二：研究人员项目与成果汇总视图', level=3)

body(
    '视图名称：v_researcher_summary。功能说明：汇总每位研究人员负责的项目总数、在研'
    '项目数（status=\'执行中\'）以及登记的成果总数，支撑工作量考核与绩效评估。'
)

code_block(
    'CREATE VIEW v_researcher_summary AS\n'
    'SELECT\n'
    '    r.researcher_id,\n'
    '    r.name AS researcher_name,\n'
    '    r.college_id,\n'
    '    COUNT(DISTINCT p.project_id) AS total_projects,\n'
    '    COUNT(DISTINCT CASE WHEN p.status = \'执行中\' THEN p.project_id END) AS active_projects,\n'
    '    COUNT(DISTINCT a.ach_id) AS total_achievements\n'
    'FROM researcher r\n'
    'LEFT JOIN project p ON r.researcher_id = p.leader_id\n'
    'LEFT JOIN achievement a ON p.project_id = a.project_id\n'
    'GROUP BY r.researcher_id, r.name, r.college_id;'
)

# View summary table
make_table(
    ['视图名称', '功能说明', '涉及基表', '关键字段', '应用场景'],
    [
        ['v_project_budget_status', '项目预算执行状态汇总', 'project, expenditure', 'budget_total, total_spent, execution_rate', '财务监控/领导看板'],
        ['v_researcher_summary', '研究人员工作量汇总', 'researcher, project, achievement', 'total_projects, active_projects, total_achievements', '工作量考核/绩效评估'],
    ]
)


# ===================== 4.4 索引设计 =====================
doc.add_heading('4.4  索引设计', level=2)

body(
    '索引（Index）是提升数据库查询性能的核心手段。合理的索引设计能够将全表扫描优化为索引'
    '扫描，极大降低I/O开销。索引设计应遵循"高频查询字段优先、外键字段必建、复合索引'
    '前缀匹配"的原则。表4-4列出了本系统设计的全部索引及其设计理由。'
)

make_wide_table(
    ['索引名', '所属表', '索引字段', '索引类型', '设计理由'],
    [
        ['idx_project_leader', 'project', 'leader_id', '普通索引',
         '高频查询：按负责人查询其负责的所有项目（"我的项目"功能），WHERE leader_id = ? 是该功能最频繁的查询条件'],
        ['idx_project_status', 'project', 'status', '普通索引',
         '高频查询：按项目状态筛选（待审/执行中/已验收等），各角色工作台均需按状态分类统计项目数量'],
        ['idx_project_end_date', 'project', 'end_date', '普通索引',
         '到期预警：查询30天内即将到期的项目（WHERE end_date BETWEEN CURDATE() AND DATE_ADD(CURDATE(), INTERVAL 30 DAY)），需按日期范围快速筛选'],
        ['idx_expenditure_project', 'expenditure', 'project_id', '普通索引',
         '高频查询：按项目查询经费流水（"经费状态"功能），每条经费记录均关联project_id，外键索引加速JOIN操作'],
        ['idx_expenditure_status', 'expenditure', 'approval_status', '普通索引',
         '待办查询：财务处查询待审批经费（WHERE approval_status = \'待审批\'），高频过滤条件'],
        ['idx_review_expert', 'review', 'expert_id', '普通索引',
         '待办查询：专家查询分配给自己待评审的项目，WHERE expert_id = ? AND result IS NULL'],
        ['idx_achievement_project', 'achievement', 'project_id', '普通索引',
         '关联查询：按项目查询成果列表，WHERE project_id = ? ORDER BY publish_date DESC'],
        ['idx_researcher_college', 'researcher', 'college_id', '普通索引',
         '分组统计：按学院分组统计研究人员，常用于学院维度统计报表的GROUP BY college_id'],
        ['idx_budget_project', 'budget', 'project_id', '普通索引',
         '关联查询：按项目查询预算科目明细，外键索引加速JOIN操作'],
    ]
)

body(
    '上述索引共计9个，覆盖了系统中所有高频查询路径。在建立索引时需注意以下几点：'
    '（1）复合索引的最左前缀原则——若某查询同时涉及project_id和approval_status，可创建'
    '(project_id, approval_status)复合索引以提高效率，但考虑到查询场景各异，本设计对单列'
    '索引与复合索引的取舍将在物理设计和实施阶段根据查询频率进一步优化；'
    '（2）索引并非越多越好——每增加一个索引都会带来额外的写入开销（INSERT/UPDATE/DELETE'
    '时须同步维护索引），因此仅对确实高频且数据量大的查询字段建立索引；'
    '（3）主键和外键默认具有索引，上表不再重复列出由PRIMARY KEY和FOREIGN KEY约束自动'
    '创建的索引。'
)


# ===================== 4.5 存储过程与函数设计 =====================
doc.add_heading('4.5  存储过程与函数设计', level=2)

body(
    '存储过程（Stored Procedure）和自定义函数（User-Defined Function）是数据库端业务逻辑'
    '封装的重要手段。存储过程适合执行复杂的多步骤业务操作（如批量更新、统计汇总），自定义'
    '函数适合嵌入SQL语句中的计算逻辑（如格式转换、等级判定）。本系统设计了两个存储过程和'
    '一个自定义函数，覆盖统计分析和业务规则计算两个核心场景。'
)

doc.add_heading('4.5.1  存储过程一：学院年度项目统计', level=3)

body(
    '存储过程名称：sp_college_project_stats(IN p_year INT)。功能说明：按指定年份统计各学院'
    '的项目申报数、立项数和立项率（立项数/申报数×100%），输出结果集供科研处统计报表模块调用。'
    '输入参数为统计年度，无输出参数，结果以SELECT结果集返回。'
)

code_block(
    'DELIMITER $$\n'
    'CREATE PROCEDURE sp_college_project_stats(IN p_year INT)\n'
    'BEGIN\n'
    '    SELECT\n'
    '        c.college_id,\n'
    '        c.name AS college_name,\n'
    '        COUNT(DISTINCT p.project_id) AS apply_count,\n'
    '        COUNT(DISTINCT CASE WHEN p.status IN (\'已立项\',\'执行中\',\'验收申请\',\'验收评审\',\'验收通过\')\n'
    '            THEN p.project_id END) AS approved_count,\n'
    '        ROUND(\n'
    '            COUNT(DISTINCT CASE WHEN p.status IN (\'已立项\',\'执行中\',\'验收申请\',\'验收评审\',\'验收通过\')\n'
    '                THEN p.project_id END) * 100.0 / NULLIF(COUNT(DISTINCT p.project_id), 0),\n'
    '            2\n'
    '        ) AS approval_rate\n'
    '    FROM college c\n'
    '    LEFT JOIN researcher r ON c.college_id = r.college_id\n'
    '    LEFT JOIN project p ON r.researcher_id = p.leader_id\n'
    '        AND YEAR(p.apply_date) = p_year\n'
    '    GROUP BY c.college_id, c.name\n'
    '    ORDER BY apply_count DESC;\n'
    'END$$\n'
    'DELIMITER ;'
)

doc.add_heading('4.5.2  存储过程二：研究人员工作量统计', level=3)

body(
    '存储过程名称：sp_researcher_workload(IN p_year INT)。功能说明：按指定年份统计每位'
    '科研人员负责的项目总数、在研项目数和成果数量，为年度工作量考核提供数据支撑。'
)

code_block(
    'DELIMITER $$\n'
    'CREATE PROCEDURE sp_researcher_workload(IN p_year INT)\n'
    'BEGIN\n'
    '    SELECT\n'
    '        r.researcher_id,\n'
    '        r.name,\n'
    '        c.name AS college_name,\n'
    '        COUNT(DISTINCT p.project_id) AS total_projects,\n'
    '        COUNT(DISTINCT CASE WHEN p.status = \'执行中\' THEN p.project_id END) AS active_projects,\n'
    '        COUNT(DISTINCT a.ach_id) AS achievement_count\n'
    '    FROM researcher r\n'
    '    LEFT JOIN college c ON r.college_id = c.college_id\n'
    '    LEFT JOIN project p ON r.researcher_id = p.leader_id\n'
    '        AND YEAR(p.apply_date) <= p_year\n'
    '    LEFT JOIN achievement a ON p.project_id = a.project_id\n'
    '    WHERE r.role = \'科研人员\'\n'
    '    GROUP BY r.researcher_id, r.name, c.name\n'
    '    ORDER BY total_projects DESC;\n'
    'END$$\n'
    'DELIMITER ;'
)

doc.add_heading('4.5.3  自定义函数：预算预警等级计算', level=3)

body(
    '函数名称：fn_budget_warning_level(p_budget_id INT)。功能说明：根据指定预算科目的已支出'
    '金额与预算金额的比值，返回三级预警等级：执行率<85%返回"正常"，≥85%且<100%返回"预警"，'
    '≥100%返回"超支"。该函数可嵌入SELECT语句中，在前端预算执行监控页面直接调用，实现'
    '实时预警标识。'
)

code_block(
    'DELIMITER $$\n'
    'CREATE FUNCTION fn_budget_warning_level(p_budget_id INT)\n'
    'RETURNS VARCHAR(10)\n'
    'DETERMINISTIC\n'
    'READS SQL DATA\n'
    'BEGIN\n'
    '    DECLARE v_amount DECIMAL(12,2);\n'
    '    DECLARE v_spent DECIMAL(12,2);\n'
    '    DECLARE v_rate DECIMAL(5,2);\n'
    '    DECLARE v_level VARCHAR(10);\n'
    '\n'
    '    SELECT amount, spent INTO v_amount, v_spent\n'
    '    FROM budget WHERE budget_id = p_budget_id;\n'
    '\n'
    '    IF v_amount IS NULL THEN\n'
    '        RETURN \'未知\';\n'
    '    END IF;\n'
    '\n'
    '    SET v_rate = (v_spent / v_amount) * 100;\n'
    '\n'
    '    IF v_rate >= 100 THEN\n'
    '        SET v_level = \'超支\';\n'
    '    ELSEIF v_rate >= 85 THEN\n'
    '        SET v_level = \'预警\';\n'
    '    ELSE\n'
    '        SET v_level = \'正常\';\n'
    '    END IF;\n'
    '\n'
    '    RETURN v_level;\n'
    'END$$\n'
    'DELIMITER ;'
)

make_table(
    ['名称', '类型', '功能说明', '输入参数', '返回/输出', '应用场景'],
    [
        ['sp_college_project_stats', '存储过程', '各学院年度项目申报数、立项数、立项率统计', 'p_year INT(年份)', 'SELECT结果集', '科研处统计报表——学院维度分析'],
        ['sp_researcher_workload', '存储过程', '各科研人员项目与成果工作量统计', 'p_year INT(年份)', 'SELECT结果集', '科研处工作量考核——人员维度分析'],
        ['fn_budget_warning_level', '自定义函数', '根据预算执行率返回三级预警等级', 'p_budget_id INT', 'VARCHAR(10) 预警等级', '经费管理模块——预算科目预警标识'],
    ]
)


# ===================== 4.6 触发器设计 =====================
doc.add_heading('4.6  触发器设计', level=2)

body(
    '触发器（Trigger）是数据库端保障数据一致性与完整性的重要机制。它能在INSERT、UPDATE、'
    'DELETE操作发生前（BEFORE）或发生后（AFTER）自动执行预定义的业务逻辑，且与触发操作'
    '在同一事务中，保证原子性。本系统设计两个核心触发器，分别服务于验收日期的自动记录和'
    '经费支出的预算一致性维护。'
)

doc.add_heading('4.6.1  触发器一：验收通过自动记录日期', level=3)

body(
    '触发器名称：trg_acceptance_update。触发时机：AFTER UPDATE ON acceptance。触发条件：'
    '当验收记录的评审结果（review_result）更新为"通过"时，自动将验收日期（acceptance_date）'
    '设置为当前日期，并将关联项目（project）的状态（status）同步更新为"验收通过"。'
    '该触发器确保验收日期不会因人工遗忘而遗漏，同时保持项目状态与验收记录的一致。'
)

code_block(
    'DELIMITER $$\n'
    'CREATE TRIGGER trg_acceptance_update\n'
    'AFTER UPDATE ON acceptance\n'
    'FOR EACH ROW\n'
    'BEGIN\n'
    '    -- 当验收评审结果更新为"通过"时\n'
    '    IF NEW.review_result = \'通过\' AND (OLD.review_result IS NULL OR OLD.review_result != \'通过\') THEN\n'
    '        -- 自动记录验收日期\n'
    '        UPDATE acceptance SET acceptance_date = CURDATE() WHERE acceptance_id = NEW.acceptance_id;\n'
    '        -- 同步更新项目状态为"验收通过"\n'
    '        UPDATE project SET status = \'验收通过\' WHERE project_id = NEW.project_id;\n'
    '    END IF;\n'
    'END$$\n'
    'DELIMITER ;'
)

doc.add_heading('4.6.2  触发器二：支出审批后自动更新预算', level=3)

body(
    '触发器名称：trg_expenditure_after_approve。触发时机：AFTER UPDATE ON expenditure。'
    '触发条件：当经费流水的审批状态（approval_status）更新为"已通过"且交易类型为"支出"时，'
    '自动将对应预算科目（budget）的已支出金额（spent）累加本次支出金额。若累加后已支出金额'
    '超过预算金额，则通过SIGNAL SQLSTATE机制回滚事务，在数据库底层拦截超支操作。'
    '该触发器是经费管理数据一致性的核心保障。'
)

code_block(
    'DELIMITER $$\n'
    'CREATE TRIGGER trg_expenditure_after_approve\n'
    'AFTER UPDATE ON expenditure\n'
    'FOR EACH ROW\n'
    'BEGIN\n'
    '    DECLARE v_budget_amount DECIMAL(12,2);\n'
    '    DECLARE v_current_spent DECIMAL(12,2);\n'
    '    DECLARE v_new_spent DECIMAL(12,2);\n'
    '\n'
    '    -- 仅当审批状态由非"已通过"变为"已通过"，且为支出类型时触发\n'
    '    IF NEW.approval_status = \'已通过\'\n'
    '       AND OLD.approval_status != \'已通过\'\n'
    '       AND NEW.type = \'支出\' THEN\n'
    '\n'
    '        -- 获取当前预算信息\n'
    '        SELECT amount, spent INTO v_budget_amount, v_current_spent\n'
    '        FROM budget WHERE budget_id = NEW.budget_id;\n'
    '\n'
    '        SET v_new_spent = v_current_spent + NEW.amount;\n'
    '\n'
    '        -- 超支拦截：若累加后超出预算，回滚事务\n'
    '        IF v_new_spent > v_budget_amount THEN\n'
    '            SIGNAL SQLSTATE \'45000\'\n'
    '            SET MESSAGE_TEXT = \'经费支出超出预算余额，无法通过审批！\';\n'
    '        END IF;\n'
    '\n'
    '        -- 更新预算已支出金额\n'
    '        UPDATE budget SET spent = v_new_spent WHERE budget_id = NEW.budget_id;\n'
    '    END IF;\n'
    'END$$\n'
    'DELIMITER ;'
)

make_table(
    ['触发器名称', '触发表', '触发时机', '触发事件', '功能说明'],
    [
        ['trg_acceptance_update', 'acceptance', 'AFTER', 'UPDATE',
         '验收评审通过时，自动记录验收日期并同步更新项目状态为"验收通过"'],
        ['trg_expenditure_after_approve', 'expenditure', 'AFTER', 'UPDATE',
         '支出审批通过时，自动更新预算已支出金额；超出预算余额时通过SIGNAL拦截并回滚'],
    ]
)

body(
    '以上两个触发器与第4.5节定义的存储过程和函数，以及第4.3节定义的视图、第4.4节定义的'
    '索引，共同构成了数据库逻辑结构设计的完整技术方案。它们在第5章物理结构设计和第6章'
    '数据库实施中将被具体化为完整的SQL DDL代码，并在第8章系统测试中接受功能与性能验证。'
)


# ---- Save ----
output_path = 'Paper/第4章_逻辑结构设计_v1.docx'
doc.save(output_path)
print(f"\n{'='*60}")
print("[DONE] Chapter 4 v1 generated successfully")
print(f"File: {output_path}")
print(f"{'='*60}")
