# -*- coding: utf-8 -*-
"""生成第2章 需求分析 DOCX — 精修版
包含: 用户需求/功能需求/DFD图/数据字典/非功能性需求
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, sys

# Ensure correct encoding for print
sys.stdout.reconfigure(encoding='utf-8', errors='replace')

doc = Document()

# ---- Page setup ----
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ---- Style setup ----
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Cm(0.74)

for lvl, sz, name in [(1, 16, 'Heading 1'), (2, 14, 'Heading 2'), (3, 13, 'Heading 3')]:
    s = doc.styles[name]
    s.font.name = '黑体'
    s.font.size = Pt(sz)
    s.font.bold = True
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if lvl == 1 else WD_ALIGN_PARAGRAPH.LEFT
    s.paragraph_format.space_before = Pt(12 if lvl == 1 else 8)
    s.paragraph_format.space_after = Pt(6 if lvl == 1 else 4)


def body(text):
    """Add body paragraph."""
    p = doc.add_paragraph(style='Normal')
    p.clear()
    r = p.add_run(text)
    r.font.name = '宋体'; r.font.size = Pt(12)
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_figure(img_path, caption, width=5.5):
    """Add centered figure with caption."""
    if not os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.add_run(f"( {caption} )").font.size = Pt(10)
        return
    pi = doc.add_paragraph()
    pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pi.paragraph_format.first_line_indent = Cm(0)
    pi.paragraph_format.space_before = Pt(6)
    pi.add_run().add_picture(img_path, width=Inches(width))
    pc = doc.add_paragraph()
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pc.paragraph_format.first_line_indent = Cm(0)
    pc.paragraph_format.space_before = Pt(2)
    pc.paragraph_format.space_after = Pt(10)
    r = pc.add_run(caption)
    r.font.name = '宋体'; r.font.size = Pt(10)
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def make_table(headers, rows, col_widths=None):
    """Create a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.name = '宋体'; r.font.size = Pt(10); r.font.bold = True
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        from docx.oxml.ns import qn as q
        cell._tc.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        shading.set(q('w:fill'), '4472C4')
        shading.set(q('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)
        r.font.color.rgb = RGBColor(255, 255, 255)
    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j < len(row) - 1 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            r.font.name = '宋体'; r.font.size = Pt(9)
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if i % 2 == 0:
                from docx.oxml.ns import qn as q
                shading = OxmlElement('w:shd')
                shading.set(q('w:fill'), 'F2F6FC')
                shading.set(q('w:val'), 'clear')
                cell._tc.get_or_add_tcPr().append(shading)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # spacer
    return table


# ================================================================
#  MATPLOTLIB SETUP
# ================================================================
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
plt.rcParams['font.sans-serif'] = ['SimHei']
plt.rcParams['font.family'] = 'sans-serif'
plt.rcParams['axes.unicode_minus'] = False


def save_fig(fig, path):
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close(fig)


# ================================================================
#  CHAPTER 2
# ================================================================
doc.add_heading('第2章  需求分析', level=1)

# ===================== 2.1 用户需求分析 =====================
doc.add_heading('2.1  用户需求分析', level=2)

body(
    '高校科研项目管理系统面向四类核心用户群体，每类用户具有差异化的业务职责和功能需求。'
    '本节逐一分析各角色的核心需求与使用场景。'
)

body('（1）科研人员。科研人员是系统的主要使用者，其核心需求包括：浏览科研处发布的'
     '项目申报通知，在线填报项目申请书并上传立项报告等附件材料，编制项目预算科目明细；'
     '随时查看本人负责项目的审批进度与当前状态；在项目执行阶段，按期提交年度进展报告，'
     '根据需要提交预算调整、延期或成员变更申请；项目到期前，提交验收申请并上传结题材料；'
     '登记项目产出的论文、专利、软件著作权等科研成果；选择对应的预算科目提交经费报销申请。')

body('（2）科研处管理员。科研处管理员是系统的核心管理者，其需求涵盖项目全生命周期：'
     '发布项目申报通知公告；对申报项目进行形式审查（审核材料完整性与规范性），通过后进入'
     '专家评审环节；为每个通过形式审查的项目分配多位评审专家，系统需防止重复分配；'
     '追踪各项目的专家评审进度，所有专家评审完成后自动进入立项公示；对立项公示项目进行'
     '最终审批；审核科研人员提交的进展报告和变更申请；组织验收评审并发放结题证书；'
     '审核科研人员登记的科研成果；按学院、年度等维度生成统计报表，支撑管理决策。')

body('（3）评审专家。评审专家登录系统后，可查看分配给自己的待评审项目列表；对每个项目'
     '提交评审意见，包括评审结果（通过/不通过）、量化评分（0-100分）和文字评语；'
     '可在评审截止前修改或撤回本人的评审意见；查看所有已评审项目的历史记录。')

body('（4）财务处管理员。财务处管理员负责科研经费的全流程监管：登记各项目的经费到账信息；'
     '审批科研人员提交的支出报销申请，审批时系统自动校验预算余额是否充足；实时查看各项目'
     '预算科目的执行率（已支出/预算总额），对预警科目（执行率≥85%）和超支科目（执行率≥100%）'
     '进行重点监控；对验收通过项目的结余经费进行处理（转入下一年或上缴学校）。')


# ---- Figure 2-1: User Role Diagram ----
try:
    fig, ax = plt.subplots(figsize=(9, 4.5))
    ax.set_xlim(0, 9); ax.set_ylim(0, 4.5); ax.axis('off')
    ax.set_facecolor('#FAFBFC'); fig.patch.set_facecolor('#FAFBFC')

    ax.text(4.5, 4.2, '高校科研项目管理系统 — 用户角色与核心需求', ha='center', va='center',
            fontsize=13, fontweight='bold', color='#1a1a2e')

    roles_data = [
        ('科研人员', 0.3, ['项目申报', '进展报告', '验收申请', '成果登记', '经费报销'], '#1565C0'),
        ('科研处管理员', 2.3, ['通知发布', '形式审查', '专家分配', '立项审批', '验收评审', '统计报表'], '#2E7D32'),
        ('评审专家', 4.3, ['待审列表', '提交评审', '打分评语', '评审记录'], '#E65100'),
        ('财务处管理员', 6.3, ['到账登记', '支出审批', '预算监控', '结余处理'], '#C62828'),
    ]

    for name, x, items, color in roles_data:
        # Role box
        rect = mpatches.FancyBboxPatch((x, 2.8), 1.8, 0.9, boxstyle="round,pad=0.06",
                                        facecolor=color, edgecolor='white', linewidth=1.5)
        ax.add_patch(rect)
        ax.text(x + 0.9, 3.25, name, ha='center', va='center', fontsize=10,
                fontweight='bold', color='white')
        # Items
        for j, item in enumerate(items):
            iy = 2.35 - j * 0.32
            ax.text(x + 0.9, iy, f'• {item}', ha='center', va='center', fontsize=7,
                    color='#333333')

    # Central system box
    rect = mpatches.FancyBboxPatch((1.5, 0.4), 6.0, 0.6, boxstyle="round,pad=0.05",
                                    facecolor='#ECEFF1', edgecolor='#546E7A', linewidth=1.2)
    ax.add_patch(rect)
    ax.text(4.5, 0.7, '高校科研项目管理系统', ha='center', va='center', fontsize=11,
            fontweight='bold', color='#1a1a2e')

    plt.tight_layout(pad=0.5)
    save_fig(fig, 'Paper/Picture/fig2_1_用户角色图.png')
    add_figure('Paper/Picture/fig2_1_用户角色图.png', '图2-1  系统用户角色与核心需求示意图')
except Exception as e:
    print(f"WARN fig2-1: {e}")


# ===================== 2.2 功能需求分析 =====================
doc.add_heading('2.2  功能需求分析', level=2)

body('系统功能按业务领域划分为八大模块，以下逐一进行详细分析。')

# 2.2.1
doc.add_heading('2.2.1  通知公告模块', level=3)
body('科研处管理员可通过该模块发布项目申报通知，填写通知标题、正文内容和发布日期，'
     '系统自动记录发布人信息。所有用户登录后可在工作台首页按发布时间倒序查看通知列表。'
     '该模块是项目申报流程的起点，确保科研人员及时获取申报信息。')

# 2.2.2
doc.add_heading('2.2.2  项目申报与审批模块', level=3)
body('该模块是系统的核心业务流程模块，覆盖项目从申报到立项的全过程。科研人员在线填写'
     '项目申报书（项目名称、项目类型（纵向/横向）、级别、负责人、起止日期、总预算等），'
     '可添加多个预算科目（科目名称及金额），并上传附件材料。提交后项目状态自动设为'
     '"申报中"。科研处管理员执行形式审查，审核材料完整性与规范性，审查通过后状态变更为'
     '"形式审查通过"。随后科研处为项目分配多位评审专家，专家提交评审意见后，系统自动'
     '检测：全部专家通过则状态自动流转为"立项公示"；任一专家不通过则退回"申报中"。'
     '科研处对立项公示项目进行最终审批，通过后状态变更为"已立项"，项目进入执行阶段。')

# 2.2.3
doc.add_heading('2.2.3  项目执行管理模块', level=3)
body('项目进入执行阶段后，科研人员可按年度提交进展报告（报告年度、提交日期、报告内容），'
     '初始状态为"待审核"，科研处审核后更新状态。科研人员可提交重要事项变更申请，涵盖'
     '三种变更类型：预算调整（关联具体预算科目及原值/新值）、延期申请（原截止日期/新截止日期）、'
     '成员变更（变更内容说明），提交后由科研处审批。提供合同/任务书上传与管理功能，'
     '合同与项目一对一关联。预算执行监控功能以表格形式展示各预算科目的预算金额、已支出金额、'
     '执行率百分比，并通过颜色标识预警状态（绿色正常、黄色预警、红色超支）。')

# 2.2.4
doc.add_heading('2.2.4  项目验收管理模块', level=3)
body('科研人员在项目到期前提交验收申请，上传结题材料附件，项目状态变更为"验收申请"。'
     '科研处管理员进行验收评审，填写评审意见和评审结果。评审通过后，项目状态变更为'
     '"验收通过"，系统触发器自动记录验收日期；评审不通过则退回"执行中"状态。'
     '对验收通过的项目，科研处可发放结题证书并上传证书文件链接。')

# 2.2.5
doc.add_heading('2.2.5  经费管理模块', level=3)
body('该模块由财务处管理员主导，科研人员协同使用。财务处登记项目经费到账信息（项目编号、'
     '到账金额、到账日期、用途说明），交易类型标记为"到账"。科研人员提交支出报销申请，'
     '需关联具体预算科目，填写金额、日期和用途，提交后状态为"待审批"。财务处进行审批时，'
     '数据库触发器自动校验：若支出金额超出对应预算科目的可用余额（预算金额-已支出金额），'
     '则拦截操作并报错；校验通过则自动更新预算表的已支出金额。系统提供项目经费执行状态'
     '总览，包括总预算、总支出、综合执行率和各科目明细。对验收通过的项目，财务处可选择'
     '结余经费处理方式（转入下一年/上缴学校）。')

# 2.2.6
doc.add_heading('2.2.6  成果管理模块', level=3)
body('科研人员可登记项目产出的科研成果，成果类型包括论文、专利、软件著作权、获奖、标准、'
     '成果转化六种，需填写标题、发表日期、作者/完成人等信息并上传附件链接。科研人员可按'
     '项目查询该项目下已登记的全部成果。科研处管理员可对提交的成果进行审核。')

# 2.2.7
doc.add_heading('2.2.7  统计报表模块', level=3)
body('该模块为科研处提供多维度的数据统计功能。按年度统计各学院的项目申报数、立项数、'
     '立项率（立项数/申报数×100%）和到账经费总额，通过数据库存储过程实现。按年度统计'
     '每位科研人员负责的项目总数、在研项目数和成果数量，为人事考核提供数据支撑。'
     '按成果类型（论文/专利/软著/获奖等）分组统计总数量，以图表形式直观展示分布比例。')

# 2.2.8
doc.add_heading('2.2.8  综合查询与工作台模块', level=3)
body('工作台（Dashboard）为各角色提供个性化的首页视图：通知公告列表（按发布时间倒序）、'
     '即将到期项目提醒（30天内到期且未验收的项目，按剩余天数升序）、各类待办事项的数量'
     '统计卡片以及快捷功能入口。综合查询模块支持按项目编号、名称、负责人、状态等条件'
     '筛选项目列表，点击可查看项目完整详情，包括基本信息、预算执行、评审记录、成果列表、'
     '经费流水和项目时间线。')

# ---- Figure 2-2: Functional Module Structure ----
try:
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 5); ax.axis('off')
    ax.set_facecolor('#FAFBFC'); fig.patch.set_facecolor('#FAFBFC')

    ax.text(5, 4.7, '系统功能模块结构图', ha='center', va='center',
            fontsize=13, fontweight='bold', color='#1a1a2e')

    modules = [
        (0.3, 2.0, 2.2, 2.2, '#E3F2FD', '通知公告\n模块', '发布通知\n列表查看'),
        (2.7, 2.0, 2.2, 2.2, '#E8F5E9', '项目申报与\n审批模块', '申报/审查\n专家分配/立项'),
        (5.1, 2.0, 2.2, 2.2, '#FFF3E0', '项目执行\n管理模块', '进展报告\n变更/合同/预算'),
        (7.5, 2.0, 2.2, 2.2, '#FCE4EC', '项目验收\n管理模块', '验收申请\n评审/证书'),
        (0.3, 0.3, 2.2, 1.5, '#F3E5F5', '经费管理\n模块', '到账/报销/审批\n预算监控/结余'),
        (2.7, 0.3, 2.2, 1.5, '#E0F2F1', '成果管理\n模块', '成果登记\n查询/审核'),
        (5.1, 0.3, 2.2, 1.5, '#FFF8E1', '统计报表\n模块', '学院/工作量\n成果分布'),
        (7.5, 0.3, 2.2, 1.5, '#ECEFF1', '工作台与\n综合查询', '工作台/待办\n项目搜索/详情'),
    ]

    for x, y, w, h, color, title_text, desc_text in modules:
        rect = mpatches.FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.06",
                                        facecolor=color, edgecolor='#90A4AE', linewidth=1.0)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2 + 0.25, title_text, ha='center', va='center',
                fontsize=9, fontweight='bold', color='#37474F')
        ax.text(x + w/2, y + h/2 - 0.35, desc_text, ha='center', va='center',
                fontsize=7.5, color='#546E7A')

    plt.tight_layout(pad=0.5)
    save_fig(fig, 'Paper/Picture/fig2_2_功能模块结构图.png')
    add_figure('Paper/Picture/fig2_2_功能模块结构图.png', '图2-2  系统功能模块结构图')
except Exception as e:
    print(f"WARN fig2-2: {e}")


# ===================== 2.3 数据流图（DFD） =====================
doc.add_heading('2.3  数据流图（DFD）', level=2)

body('数据流图（Data Flow Diagram）是结构化需求分析的核心工具。本节采用"自顶向下、'
     '逐层细化"的策略，首先绘制顶层数据流图刻画系统与外部实体之间的数据交换全貌，'
     '然后绘制一层数据流图分解系统的主要处理功能及其间的数据流动，最后选取项目申报'
     '子功能绘制二层数据流图进行精细化描述。')

# -- 2.3.1 顶层DFD --
doc.add_heading('2.3.1  顶层数据流图', level=3)

body('顶层数据流图（上下文图）将整个系统视为单一处理节点，重点展示系统与四类外部实体'
     '（科研人员、科研处管理员、评审专家、财务处管理员）之间的数据交互关系。科研人员'
     '向系统提交申报信息、报告信息、验收申请、成果信息和报销申请，并从系统获取通知信息、'
     '审批反馈和项目状态。科研处管理员向系统发布通知、提交审查意见、分配评审专家、进行'
     '审批决策，并从系统获取申报材料、评审结果和统计数据。评审专家从系统获取待审项目'
     '信息，向系统提交评审意见。财务处管理员向系统登记到账信息、提交审批决策，并从系统'
     '获取经费状态和预算执行情况。')

# -- Top-level DFD figure --
try:
    fig, ax = plt.subplots(figsize=(10, 5.5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 5.5); ax.axis('off')
    ax.set_facecolor('white'); fig.patch.set_facecolor('white')

    # Title
    ax.text(5, 5.25, '顶层数据流图（上下文图）', ha='center', va='center',
            fontsize=12, fontweight='bold', color='#1a1a2e')

    # Central system process
    circ = mpatches.FancyBboxPatch((2.8, 2.1), 4.4, 2.0, boxstyle="round,pad=0.1",
                                    facecolor='#E3F2FD', edgecolor='#1565C0', linewidth=2)
    ax.add_patch(circ)
    ax.text(5, 3.1, '高校科研项目\n管理系统', ha='center', va='center',
            fontsize=12, fontweight='bold', color='#0D47A1')

    # External entities
    entities = [
        (0.2, 3.5, '科研人员', '#E8F5E9', '#2E7D32'),
        (7.8, 3.8, '科研处\n管理员', '#FFF3E0', '#E65100'),
        (8.0, 1.8, '评审专家', '#FCE4EC', '#C62828'),
        (0.2, 1.5, '财务处\n管理员', '#F3E5F5', '#6A1B9A'),
    ]
    for x, y, name, fc, ec in entities:
        rect = mpatches.FancyBboxPatch((x, y), 1.8, 1.0, boxstyle="round,pad=0.05",
                                        facecolor=fc, edgecolor=ec, linewidth=1.5)
        ax.add_patch(rect)
        ax.text(x + 0.9, y + 0.5, name, ha='center', va='center',
                fontsize=9, fontweight='bold', color=ec)

    # Data flow labels
    flows = [
        (0.5, 4.8, 5.0, 4.3, '申报、报告、验收、成果、报销', '#2E7D32'),
        (5.0, 4.5, 0.5, 4.5, '通知、审批结果、项目状态', '#1565C0'),
        (7.8, 3.3, 5.0, 3.8, '通知、审查、分配、审批', '#E65100'),
        (5.0, 3.5, 7.8, 3.0, '申报材料、评审结果、统计数据', '#1565C0'),
        (8.2, 2.5, 5.0, 2.5, '评审意见、打分', '#C62828'),
        (5.0, 2.8, 7.8, 2.2, '待审项目信息', '#1565C0'),
        (0.5, 2.0, 5.0, 2.2, '到账登记、审批决策', '#6A1B9A'),
        (5.0, 2.0, 0.5, 1.8, '经费状态、预算执行情况', '#1565C0'),
    ]
    for x1, y1, x2, y2, label, color in flows:
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color=color, lw=1.0))
        ax.text((x1 + x2) / 2, (y1 + y2) / 2 + 0.08, label, ha='center', va='center',
                fontsize=6.5, color=color, style='italic')

    plt.tight_layout(pad=0.3)
    save_fig(fig, 'Paper/Picture/fig2_3_顶层DFD.png')
    add_figure('Paper/Picture/fig2_3_顶层DFD.png', '图2-3  顶层数据流图（上下文图）')
except Exception as e:
    print(f"WARN fig2-3: {e}")
    body('（图2-3：顶层数据流图——待补充）')


# -- 2.3.2 一层DFD --
doc.add_heading('2.3.2  一层数据流图', level=3)

body('一层数据流图将顶层图中的单一处理节点分解为八个核心处理功能，展示了各功能模块之间'
     '以及功能模块与数据存储之间的数据流动关系。八个处理功能分别为：P1-通知管理、'
     'P2-项目申报与审批、P3-项目执行管理、P4-项目验收管理、P5-经费管理、P6-成果管理、'
     'P7-统计报表、P8-综合查询。数据存储包括：D1-通知表、D2-项目表、D3-预算表、'
     'D4-评审记录表、D5-经费流水表、D6-成果表、D7-进展报告表、D8-验收记录表等。')

# -- Level-1 DFD figure --
try:
    fig, ax = plt.subplots(figsize=(11, 7))
    ax.set_xlim(0, 11); ax.set_ylim(0, 7); ax.axis('off')
    ax.set_facecolor('white'); fig.patch.set_facecolor('white')

    ax.text(5.5, 6.75, '一层数据流图', ha='center', va='center',
            fontsize=12, fontweight='bold', color='#1a1a2e')

    # Processes (P1-P8)
    procs = [
        (0.3, 4.8, 2.3, 1.3, '#E3F2FD', 'P1\n通知管理'),
        (3.0, 5.2, 2.3, 1.3, '#E8F5E9', 'P2\n项目申报与审批'),
        (6.0, 5.2, 2.3, 1.3, '#FFF3E0', 'P3\n项目执行管理'),
        (8.5, 5.2, 2.3, 1.3, '#FCE4EC', 'P4\n项目验收管理'),
        (0.3, 2.6, 2.3, 1.3, '#F3E5F5', 'P5\n经费管理'),
        (3.0, 2.6, 2.3, 1.3, '#E0F2F1', 'P6\n成果管理'),
        (6.0, 2.6, 2.3, 1.3, '#FFF8E1', 'P7\n统计报表'),
        (8.5, 2.6, 2.3, 1.3, '#ECEFF1', 'P8\n综合查询'),
    ]
    for x, y, w, h, color, label in procs:
        rect = mpatches.FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.05",
                                        facecolor=color, edgecolor='#607D8B', linewidth=1.2)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2, label, ha='center', va='center',
                fontsize=9, fontweight='bold', color='#37474F')

    # Data stores (bottom)
    stores = ['D1 通知表', 'D2 项目表', 'D3 预算表', 'D4 评审表', 'D5 经费表', 'D6 成果表']
    for i, s in enumerate(stores):
        sx = 0.5 + i * 1.7
        rect = mpatches.FancyBboxPatch((sx, 0.4), 1.5, 0.7, boxstyle="round,pad=0.04",
                                        facecolor='#ECEFF1', edgecolor='#90A4AE', linewidth=1.0)
        ax.add_patch(rect)
        ax.text(sx + 0.75, 0.75, s, ha='center', va='center', fontsize=7.5, color='#546E7A')

    ax.annotate('', xy=(1.0, 3.8), xytext=(1.0, 1.2),
                arrowprops=dict(arrowstyle='->', color='#90A4AE', lw=0.8))
    ax.annotate('', xy=(3.5, 3.8), xytext=(3.5, 1.2),
                arrowprops=dict(arrowstyle='->', color='#90A4AE', lw=0.8))

    plt.tight_layout(pad=0.3)
    save_fig(fig, 'Paper/Picture/fig2_4_一层DFD.png')
    add_figure('Paper/Picture/fig2_4_一层DFD.png', '图2-4  一层数据流图')
except Exception as e:
    print(f"WARN fig2-4: {e}")
    body('（图2-4：一层数据流图——待补充）')


# -- 2.3.3 二层DFD --
doc.add_heading('2.3.3  二层数据流图（项目申报与审批）', level=3)

body('选取核心处理功能"P2-项目申报与审批"进行精细化分解，得到五个子处理：P2.1-申报通知'
     '查看、P2.2-项目申报提交、P2.3-形式审查、P2.4-专家评审、P2.5-立项审批。该分解完整'
     '刻画了项目从申报到立项的内部数据处理逻辑与各子功能间的数据传递关系。')

try:
    fig, ax = plt.subplots(figsize=(11, 6))
    ax.set_xlim(0, 11); ax.set_ylim(0, 6); ax.axis('off')
    ax.set_facecolor('white'); fig.patch.set_facecolor('white')

    ax.text(5.5, 5.75, '二层数据流图 —— P2 项目申报与审批', ha='center', va='center',
            fontsize=12, fontweight='bold', color='#1a1a2e')

    sub_procs = [
        (0.2, 3.6, 2.0, 1.2, '#E3F2FD', 'P2.1\n申报通知查看'),
        (2.6, 3.6, 2.0, 1.2, '#E8F5E9', 'P2.2\n项目申报提交'),
        (5.0, 3.6, 2.0, 1.2, '#FFF3E0', 'P2.3\n形式审查'),
        (7.4, 3.6, 2.0, 1.2, '#FCE4EC', 'P2.4\n专家评审'),
        (9.0, 3.6, 1.8, 1.2, '#F3E5F5', 'P2.5\n立项审批'),
    ]
    for x, y, w, h, color, label in sub_procs:
        rect = mpatches.FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.05",
                                        facecolor=color, edgecolor='#607D8B', linewidth=1.2)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2, label, ha='center', va='center',
                fontsize=8.5, fontweight='bold', color='#37474F')

    # External entities
    for x, y, name, color in [(0.2, 5.2, '科研人员', '#2E7D32'),
                               (7.0, 5.2, '科研处管理员', '#E65100'),
                               (8.5, 5.2, '评审专家', '#C62828')]:
        rect = mpatches.FancyBboxPatch((x, y), 1.8, 0.5, boxstyle="round,pad=0.04",
                                        facecolor='white', edgecolor=color, linewidth=1.5)
        ax.add_patch(rect)
        ax.text(x + 0.9, y + 0.25, name, ha='center', va='center', fontsize=8, color=color)

    # Data stores
    stores2 = ['D1 通知表', 'D2 项目表', 'D3 预算表', 'D4 评审记录表']
    for i, s in enumerate(stores2):
        sx = 0.5 + i * 2.7
        rect = mpatches.FancyBboxPatch((sx, 0.3), 2.3, 0.6, boxstyle="round,pad=0.04",
                                        facecolor='#ECEFF1', edgecolor='#90A4AE', linewidth=1.0)
        ax.add_patch(rect)
        ax.text(sx + 1.15, 0.6, s, ha='center', va='center', fontsize=7.5, color='#546E7A')

    # Flow arrows (simplified)
    for i in range(4):
        x1 = 0.2 + i * 2.4 + 1.0
        x2 = x1 + 2.2
        ax.annotate('', xy=(x2, 4.2), xytext=(x1 + 1.8, 4.2),
                    arrowprops=dict(arrowstyle='->', color='#78909C', lw=1.2))

    plt.tight_layout(pad=0.3)
    save_fig(fig, 'Paper/Picture/fig2_5_二层DFD.png')
    add_figure('Paper/Picture/fig2_5_二层DFD.png', '图2-5  二层数据流图（P2-项目申报与审批）')
except Exception as e:
    print(f"WARN fig2-5: {e}")
    body('（图2-5：二层数据流图——待补充）')


# ===================== 2.4 数据字典（DD） =====================
doc.add_heading('2.4  数据字典（DD）', level=2)

body('数据字典对数据流图中出现的所有数据元素进行精确的语义定义，是数据库概念结构设计的'
     '直接依据。以下按数据项、数据流、数据存储三类进行系统化定义。')

# 2.4.1 数据项
doc.add_heading('2.4.1  数据项定义', level=3)

body('数据项是系统中不可再分的最小数据单位。系统核心数据项定义如下表所示。')

make_table(
    ['编号', '数据项名称', '含义说明', '类型', '长度', '取值范围/格式', '备注'],
    [
        ['DI-01', 'project_id', '项目编号', '字符型', '20', '如 P2024001', '主键，系统自动生成'],
        ['DI-02', 'project_name', '项目名称', '字符型', '200', '中文/英文/数字', '必填'],
        ['DI-03', 'project_type', '项目类型', '枚举型', '-', '纵向 / 横向', '必填'],
        ['DI-04', 'project_level', '项目级别', '字符型', '20', '国家级/省部级/市厅级/横向', '—'],
        ['DI-05', 'project_status', '项目状态', '枚举型', '-', '10种状态值', '默认"申报中"'],
        ['DI-06', 'budget_total', '总预算（万元）', '数值型', '12,2', '> 0', '必填'],
        ['DI-07', 'researcher_id', '用户工号', '字符型', '10', 'R开头+4位数字', '主键'],
        ['DI-08', 'user_name', '用户姓名', '字符型', '20', '中文', '必填'],
        ['DI-09', 'user_password', '登录密码', '字符型', '100', '字母/数字/符号', '默认123456'],
        ['DI-10', 'user_role', '用户角色', '枚举型', '-', '科研人员/专家/科研处/财务处', '—'],
        ['DI-11', 'college_id', '学院编号', '字符型', '4', '如 C001', '外键→college'],
        ['DI-12', 'budget_category', '预算科目名称', '字符型', '50', '设备费/材料费/差旅费等', '必填'],
        ['DI-13', 'budget_amount', '预算金额（万元）', '数值型', '12,2', '> 0', '必填'],
        ['DI-14', 'budget_spent', '已支出金额（万元）', '数值型', '12,2', '≥ 0', '默认0'],
        ['DI-15', 'exp_type', '交易类型', '枚举型', '-', '到账 / 支出', '必填'],
        ['DI-16', 'exp_amount', '交易金额（万元）', '数值型', '12,2', '> 0', '必填'],
        ['DI-17', 'exp_approval', '审批状态', '枚举型', '-', '待审批/已通过/已驳回', '默认"待审批"'],
        ['DI-18', 'review_score', '评审打分', '整数型', '-', '0 - 100', '专家打分'],
        ['DI-19', 'review_result', '评审结果', '枚举型', '-', '通过 / 不通过', '必填'],
        ['DI-20', 'ach_type', '成果类型', '枚举型', '-', '论文/专利/软著/获奖/标准/转化', '六选一'],
        ['DI-21', 'start_date', '开始日期', '日期型', '-', 'YYYY-MM-DD', '—'],
        ['DI-22', 'end_date', '结束日期', '日期型', '-', 'YYYY-MM-DD', '—'],
        ['DI-23', 'apply_date', '申请日期', '日期型', '-', 'YYYY-MM-DD', '默认当前日期'],
        ['DI-24', 'acceptance_date', '验收日期', '日期型', '-', 'YYYY-MM-DD', '触发器自动填充'],
    ]
)

# 2.4.2 数据流
doc.add_heading('2.4.2  数据流定义', level=3)

body('数据流描述数据在系统中流动的方向与内容。系统主要数据流定义如下。')

make_table(
    ['编号', '数据流名称', '来源', '去向', '组成', '说明'],
    [
        ['DF-01', '申报通知', '科研处管理员', '科研人员', '通知标题+内容+发布日期', '科研处发布→全员查看'],
        ['DF-02', '项目申报信息', '科研人员', 'P2.2 项目申报提交', '项目编号+名称+类型+级别+预算+附件', '在线填报提交'],
        ['DF-03', '形式审查结果', 'P2.3 形式审查', 'D2 项目表', '审查结果(通过/不通过)+审查意见', '更新项目状态'],
        ['DF-04', '专家分配信息', '科研处管理员', 'D4 评审记录表', '项目编号+专家工号列表', '为项目分配评审专家'],
        ['DF-05', '评审意见', '评审专家', 'D4 评审记录表', '项目编号+专家工号+结果+打分+评语', '专家提交评审'],
        ['DF-06', '立项审批结果', 'P2.5 立项审批', 'D2 项目表', '审批结果(通过/不通过)', '更新为"已立项"'],
        ['DF-07', '进展报告', '科研人员', 'D7 进展报告表', '项目编号+年度+日期+内容', '按年度提交'],
        ['DF-08', '变更申请', '科研人员', 'D2+/变更申请表', '变更类型+原值+新值+理由', '预算/延期/成员'],
        ['DF-09', '验收申请', '科研人员', 'D8 验收记录表', '项目编号+申请日期+结题材料', '到期前提交'],
        ['DF-10', '经费到账信息', '财务处管理员', 'D5 经费流水表', '项目编号+金额+日期+用途', '类型=到账'],
        ['DF-11', '报销申请', '科研人员', 'D5 经费流水表', '项目编号+预算科目+金额+用途', '需关联预算科目'],
        ['DF-12', '审批结果', '财务处管理员', 'D5 经费流水表', '审批结果(通过/驳回)', '触发器自动更新预算'],
        ['DF-13', '成果登记信息', '科研人员', 'D6 成果表', '项目编号+类型+标题+日期+作者+附件', '登记成果'],
        ['DF-14', '统计数据', 'D2-D6各数据表', 'P7 统计报表', '汇总/分组统计结果', '存储过程生成'],
    ]
)

# 2.4.3 数据存储
doc.add_heading('2.4.3  数据存储定义', level=3)

body('数据存储描述系统中需要持久化保存的数据集合，对应数据库中的核心表。')

make_table(
    ['编号', '数据存储名称', '对应数据库表', '组成说明', '主键', '关联'],
    [
        ['DS-01', '学院信息', 'college', '学院编号、名称、院长、电话', 'college_id', '被researcher引用'],
        ['DS-02', '用户信息', 'researcher', '工号、姓名、密码、职称、学院、联系方式、角色', 'researcher_id', '引用college'],
        ['DS-03', '项目信息', 'project', '项目编号、名称、类型、级别、负责人、日期、预算、状态', 'project_id', '引用researcher'],
        ['DS-04', '预算科目', 'budget', '科目ID、项目编号、科目名称、预算金额、已支出金额', 'budget_id', '引用project'],
        ['DS-05', '经费流水', 'expenditure', '流水ID、项目编号、预算科目、类型、金额、日期、审批状态', 'exp_id', '引用project/budget'],
        ['DS-06', '评审记录', 'review', '项目编号、专家工号、评审结果、打分(0-100)、评语、评审日期', '(project_id,expert_id)', '引用project/researcher'],
        ['DS-07', '通知公告', 'notice', '通知ID、标题、内容、发布日期、发布人工号', 'notice_id', '引用researcher'],
        ['DS-08', '进展报告', 'progress_report', '报告ID、项目编号、报告年度、提交日期、内容、审核状态', 'report_id', '引用project'],
        ['DS-09', '变更申请', 'change_request', '变更ID、项目编号、变更类型、原值、新值、理由、审批状态', 'change_id', '引用project'],
        ['DS-10', '验收记录', 'acceptance', '验收ID、项目编号、申请日期、评审结果、验收日期、证书链接', 'acceptance_id', '引用project'],
        ['DS-11', '科研成果', 'achievement', '成果ID、项目编号、成果类型、标题、发表日期、作者、附件', 'ach_id', '引用project'],
        ['DS-12', '合同信息', 'contract', '合同ID、项目编号、文件路径、签署日期、内容描述', 'contract_id', '引用project'],
    ]
)


# ===================== 2.5 非功能性需求 =====================
doc.add_heading('2.5  非功能性需求', level=2)

body('非功能性需求定义了系统在功能正确性之外必须满足的质量属性与约束条件，是系统设计'
     '与实现的重要基准。以下从五个维度进行定义。')

doc.add_heading('2.5.1  性能需求', level=3)
body('（1）页面首次加载时间不超过3秒；（2）简单查询API响应时间不超过1秒，复杂统计查询'
     '不超过3秒；（3）系统应支持至少50名用户同时在线操作而不出现明显性能退化；'
     '（4）对高频查询字段（如项目编号、负责人工号、项目状态）建立数据库索引，确保百万级'
     '记录下的查询性能；（5）存储过程执行时间应在合理范围内，复杂统计过程不超过5秒。')

doc.add_heading('2.5.2  安全性需求', level=3)
body('（1）数据库密码等敏感配置不得硬编码在源代码中，须通过环境变量或外部配置文件管理；'
     '（2）所有SQL查询必须使用参数化查询（Parameterized Query），从根本上杜绝SQL注入攻击；'
     '（3）文件上传须校验文件类型（白名单）与大小限制，防范恶意文件上传；'
     '（4）用户密码应采用bcrypt等单向加密算法存储，当前系统使用明文存储，为已知待改进项；'
     '（5）生产环境须关闭Flask Debug模式，避免敏感调试信息泄露；'
     '（6）CORS跨域配置应设置明确的允许来源白名单，而非通配符"*"。')

doc.add_heading('2.5.3  可用性需求', level=3)
body('（1）用户界面采用响应式设计，适配主流分辨率（1366×768及以上）的桌面浏览器；'
     '（2）所有操作须提供明确的成功/失败反馈提示，关键操作（如删除、审批）须有二次确认机制；'
     '（3）表单须明确标识必填字段（红色星号），提供输入格式提示（如日期选择器、下拉列表）；'
     '（4）预算执行率达到85%以上的科目以黄色预警标识，执行率达到100%以上的科目以红色超支'
     '标识，帮助用户快速识别风险。')

doc.add_heading('2.5.4  可维护性需求', level=3)
body('（1）源代码须有清晰的中文注释，标注每个模块的功能、参数与返回值；'
     '（2）数据库脚本须包含完整的建库、建表、测试数据插入、存储过程、触发器及视图的创建语句，'
     '确保一键部署与重建能力；（3）后端API路由遵循RESTful规范，URL命名语义化，HTTP方法语义化；'
     '（4）所有API统一返回JSON格式 {"code": 200, "data": ..., "message": "..."}；'
     '（5）前端代码按功能模块拆分为独立的JavaScript文件，便于维护与扩展。')

doc.add_heading('2.5.5  数据完整性需求', level=3)
body('（1）所有外键关系在数据库层面通过FOREIGN KEY约束明确定义，不得仅依赖应用层维护；'
     '（2）项目删除时级联删除关联的预算记录（ON DELETE CASCADE），避免孤数据；'
     '（3）研究人员删除时其关联的学院外键设为NULL（ON DELETE SET NULL），保留数据可追溯性；'
     '（4）经费支出审批通过后，由数据库触发器自动更新对应预算科目的已支出金额，并在支出超出'
     '预算余额时通过SIGNAL SQLSTATE机制在数据库底层拦截，确保数据一致性不受应用层缺陷影响；'
     '（5）验收通过后由触发器自动记录验收日期，避免人工录入遗漏。')


# ---- Save ----
output_path = 'Paper/第2章_需求分析_v1.docx'
doc.save(output_path)
print(f"\n{'='*60}")
print("[DONE] Chapter 2 v1 generated successfully")
print(f"File: {output_path}")
print(f"{'='*60}")
