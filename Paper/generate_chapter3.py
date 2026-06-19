# -*- coding: utf-8 -*-
"""生成第3章 数据库概念结构设计 DOCX — 精修版
包含: 局部E-R图(3张) / 全局E-R图 / 冲突消解说明
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, sys

sys.stdout.reconfigure(encoding='utf-8', errors='replace')

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

style = doc.styles['Normal']
style.font.name = '宋体'; style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Cm(0.74)

for lvl, sz in [(1, 16), (2, 14), (3, 13)]:
    s = doc.styles[f'Heading {lvl}']
    s.font.name = '黑体'; s.font.size = Pt(sz); s.font.bold = True
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if lvl == 1 else WD_ALIGN_PARAGRAPH.LEFT
    s.paragraph_format.space_before = Pt(12 if lvl == 1 else 8)
    s.paragraph_format.space_after = Pt(6 if lvl == 1 else 4)


def body(text):
    p = doc.add_paragraph(style='Normal'); p.clear()
    r = p.add_run(text)
    r.font.name = '宋体'; r.font.size = Pt(12)
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_figure(img_path, caption, width=5.8):
    if not os.path.exists(img_path):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.add_run(f"( {caption} )").font.size = Pt(10)
        return
    pi = doc.add_paragraph(); pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pi.paragraph_format.first_line_indent = Cm(0)
    pi.paragraph_format.space_before = Pt(6)
    pi.add_run().add_picture(img_path, width=Inches(width))
    pc = doc.add_paragraph(); pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pc.paragraph_format.first_line_indent = Cm(0)
    pc.paragraph_format.space_before = Pt(2)
    pc.paragraph_format.space_after = Pt(10)
    r = pc.add_run(caption)
    r.font.name = '宋体'; r.font.size = Pt(10)
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def make_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    from docx.enum.table import WD_TABLE_ALIGNMENT
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.name = '宋体'; r.font.size = Pt(10); r.font.bold = True
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), '4472C4'); shd.set(qn('w:val'), 'clear')
        c._tc.get_or_add_tcPr().append(shd)
        r.font.color.rgb = RGBColor(255, 255, 255)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.rows[i + 1].cells[j]; c.text = ''
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            r.font.name = '宋体'; r.font.size = Pt(9)
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if i % 2 == 0:
                shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'F2F6FC'); shd.set(qn('w:val'), 'clear')
                c._tc.get_or_add_tcPr().append(shd)
    doc.add_paragraph()
    return t


# ================================================================
#  MATPLOTLIB SETUP
# ================================================================
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Arc, Ellipse
import numpy as np

plt.rcParams['font.sans-serif'] = ['SimHei']
plt.rcParams['font.family'] = 'sans-serif'
plt.rcParams['axes.unicode_minus'] = False


def save_fig(fig, path):
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close(fig)


def draw_entity(ax, x, y, name, attrs, pk_idx=0, width=2.4, color='#E3F2FD', edge='#1565C0'):
    """Draw entity box with attributes.
    attrs: list of (attr_name, is_key) tuples
    """
    n = len(attrs)
    height = max(1.0, 0.5 + n * 0.28)
    rect = FancyBboxPatch((x - width/2, y - height/2), width, height,
                           boxstyle="round,pad=0.08", facecolor=color,
                           edgecolor=edge, linewidth=1.5)
    ax.add_patch(rect)
    # Entity name at top
    ax.text(x, y + height/2 - 0.22, name, ha='center', va='center',
            fontsize=9, fontweight='bold', color='#0D47A1')

    # Separator line
    ax.plot([x - width/2 + 0.15, x + width/2 - 0.15],
            [y + height/2 - 0.4, y + height/2 - 0.4],
            color=edge, linewidth=0.6)

    # Attributes
    for i, attr in enumerate(attrs):
        ay = y + height/2 - 0.6 - i * 0.28
        txt = attr
        ax.text(x, ay, txt, ha='center', va='center', fontsize=7,
                color='#37474F' if i != pk_idx else '#C62828',
                fontweight='bold' if i == pk_idx else 'normal')
        if i == pk_idx:
            # Underline for PK
            tw = len(txt) * 0.08
            ax.plot([x - tw/2, x + tw/2], [ay - 0.1, ay - 0.1],
                    color='#C62828', linewidth=1.2)


def draw_relationship(ax, x, y, name, rel_type='', width=1.0, height=0.55):
    """Draw a diamond relationship symbol."""
    diamond = mpatches.Polygon([
        (x, y + height/2), (x + width/2, y), (x, y - height/2), (x - width/2, y)
    ], facecolor='#FFF9C4', edgecolor='#F57F17', linewidth=1.2)
    ax.add_patch(diamond)
    ax.text(x, y, name, ha='center', va='center', fontsize=7.5,
            fontweight='bold', color='#E65100')


def draw_cardinality(ax, x, y, text):
    """Draw cardinality label near a connection point."""
    ax.text(x, y, text, ha='center', va='center', fontsize=7,
            fontweight='bold', color='#546E7A',
            bbox=dict(boxstyle='round,pad=0.1', facecolor='white',
                      edgecolor='#CFD8DC', alpha=0.9))


# ================================================================
#  CHAPTER 3
# ================================================================
doc.add_heading('第3章  数据库概念结构设计', level=1)

body(
    '概念结构设计是数据库设计的核心阶段，其任务是将需求分析阶段获得的用户需求抽象为'
    '信息世界中的概念模型。本章采用"自底向上"的策略：首先按业务域分别设计局部E-R图，'
    '然后将各局部E-R图合并为全局E-R图，并在合并过程中消解各类冲突。E-R图（Entity-Relationship '
    'Diagram）作为概念模型的主要表达工具，直观地刻画了现实世界中的实体（Entity）、属性'
    '（Attribute）及实体间的联系（Relationship）。'
)

# ===================== 3.1 局部E-R图设计 =====================
doc.add_heading('3.1  局部E-R图设计', level=2)

body(
    '局部E-R图的设计遵循"分而治之"原则，按系统核心业务域分别建模。本节设计了三个'
    '局部E-R图，分别覆盖用户与组织管理、项目申报与评审管理、经费与成果管理三大业务域，'
    '每个局部E-R图均完整标识了该域内的实体、属性（含主键标识）、联系类型及联系的属性。'
)

# ---- 3.1.1 局部E-R图1: 用户与项目管理 ----
doc.add_heading('3.1.1  局部E-R图一：用户与项目管理', level=3)

body(
    '该局部E-R图覆盖用户管理与项目核心信息两大业务域，包含三个实体：学院（College）、'
    '研究人员（Researcher）和科研项目（Project）。学院与研究人员之间为1:n的"归属"联系'
    '（一个学院包含多名研究人员，一名研究人员仅属于一个学院）。研究人员与科研项目之间'
    '存在两个联系：1:n的"负责"联系（一名研究人员可负责多个项目，每个项目仅有一名负责人）'
    '和m:n的"参与"联系（多名研究人员可参与一个项目，一名研究人员可参与多个项目）。'
    '"参与"联系具有属性"参与角色"（如项目骨干、一般成员），用于标识研究人员在项目中'
    '承担的具体职责。'
)

try:
    fig, ax = plt.subplots(figsize=(11, 6.5))
    ax.set_xlim(0, 11); ax.set_ylim(0, 6.5); ax.axis('off')
    ax.set_facecolor('white'); fig.patch.set_facecolor('white')
    ax.text(5.5, 6.2, '局部E-R图一：用户与项目管理', ha='center', va='center',
            fontsize=12, fontweight='bold', color='#1a1a2e')

    # Entity: College
    draw_entity(ax, 1.5, 4.5, '学院 (College)', [
        'college_id (PK)', 'name', 'dean', 'tel'
    ], pk_idx=0, width=2.4, color='#E3F2FD', edge='#1565C0')

    # Entity: Researcher
    draw_entity(ax, 5.5, 4.5, '研究人员 (Researcher)', [
        'researcher_id (PK)', 'name', 'password', 'title',
        'phone', 'email', 'role'
    ], pk_idx=0, width=2.6, color='#E8F5E9', edge='#2E7D32')

    # Entity: Project
    draw_entity(ax, 5.5, 1.5, '科研项目 (Project)', [
        'project_id (PK)', 'name', 'type', 'level',
        'start_date', 'end_date', 'budget_total', 'status'
    ], pk_idx=0, width=2.6, color='#FFF3E0', edge='#E65100')

    # Relationship: 归属 (College -> Researcher, 1:n)
    draw_relationship(ax, 3.5, 4.5, '归属', width=0.9, height=0.5)
    ax.annotate('', xy=(4.2, 4.5), xytext=(3.95, 4.5),
                arrowprops=dict(arrowstyle='-', color='#37474F', lw=1.2))
    ax.annotate('', xy=(2.8, 4.5), xytext=(3.05, 4.5),
                arrowprops=dict(arrowstyle='-', color='#37474F', lw=1.2))
    draw_cardinality(ax, 2.45, 4.75, '1')
    draw_cardinality(ax, 4.55, 4.75, 'n')

    # Relationship: 负责 (Researcher -> Project, 1:n)
    draw_relationship(ax, 5.5, 3.0, '负责', width=0.9, height=0.5)
    ax.annotate('', xy=(5.5, 3.25), xytext=(5.5, 3.6),
                arrowprops=dict(arrowstyle='-', color='#37474F', lw=1.2))
    ax.annotate('', xy=(5.5, 2.25), xytext=(5.5, 2.75),
                arrowprops=dict(arrowstyle='-', color='#37474F', lw=1.2))
    draw_cardinality(ax, 5.15, 4.0, '1')
    draw_cardinality(ax, 5.15, 2.0, 'n')

    # Relationship: 参与 (Researcher <-> Project, m:n)
    draw_relationship(ax, 8.0, 3.0, '参与', width=0.8, height=0.5)
    # m side
    ax.annotate('', xy=(7.6, 3.0), xytext=(6.85, 3.6),
                arrowprops=dict(arrowstyle='-', color='#37474F', lw=1.2))
    draw_cardinality(ax, 7.0, 3.8, 'm')
    # n side
    ax.annotate('', xy=(7.6, 3.0), xytext=(6.85, 2.4),
                arrowprops=dict(arrowstyle='-', color='#37474F', lw=1.2))
    draw_cardinality(ax, 7.0, 2.2, 'n')
    # Relationship attribute
    ax.text(9.1, 3.0, '参与角色', ha='center', va='center', fontsize=7,
            color='#6A1B9A', style='italic')
    ax.annotate('', xy=(8.4, 3.0), xytext=(8.7, 3.0),
                arrowprops=dict(arrowstyle='-', color='#6A1B9A', lw=0.8))

    plt.tight_layout(pad=0.5)
    save_fig(fig, 'Paper/Picture/fig3_1_局部ER1.png')
    add_figure('Paper/Picture/fig3_1_局部ER1.png', '图3-1  局部E-R图一：用户与项目管理')
except Exception as e:
    print(f"WARN fig3-1: {e}")
    body('（图3-1：局部E-R图一——待补充）')


# ---- 3.1.2 局部E-R图2: 项目审批与评审 ----
doc.add_heading('3.1.2  局部E-R图二：项目审批与评审管理', level=3)

body(
    '该局部E-R图覆盖项目申报审批与专家评审业务域。通知公告（Notice）实体由科研处管理员'
    '（Researcher实体，role=科研处）发布，与科研项目之间存在1:n的"引导申报"联系'
    '（一份通知可引导多个项目申报）。科研项目与评审专家之间存在m:n的"评审"联系'
    '（一个项目由多位专家评审，一位专家可评审多个项目），该联系具有属性：评审结果'
    '（通过/不通过）、评审打分（0-100分）和评审评语。合同（Contract）实体与科研项目'
    '之间为1:1的"签署"联系（一个项目签署一份合同）。'
)

try:
    fig, ax = plt.subplots(figsize=(11, 6.5))
    ax.set_xlim(0, 11); ax.set_ylim(0, 6.5); ax.axis('off')
    ax.set_facecolor('white'); fig.patch.set_facecolor('white')
    ax.text(5.5, 6.2, '局部E-R图二：项目审批与评审管理', ha='center', va='center',
            fontsize=12, fontweight='bold', color='#1a1a2e')

    # Entity: Notice
    draw_entity(ax, 1.5, 4.5, '通知公告 (Notice)', [
        'notice_id (PK)', 'title', 'content',
        'publish_date', 'publisher_id'
    ], pk_idx=0, width=2.4, color='#F3E5F5', edge='#6A1B9A')

    # Entity: Researcher (simplified - just show expert side)
    draw_entity(ax, 5.5, 5.2, '研究人员 (Researcher)', [
        'researcher_id (PK)', 'name', 'title',
        'college_id', 'role'
    ], pk_idx=0, width=2.4, color='#E8F5E9', edge='#2E7D32')

    # Entity: Project
    draw_entity(ax, 5.5, 2.2, '科研项目 (Project)', [
        'project_id (PK)', 'name', 'type', 'level',
        'leader_id', 'status', 'budget_total'
    ], pk_idx=0, width=2.5, color='#FFF3E0', edge='#E65100')

    # Entity: Contract
    draw_entity(ax, 9.0, 2.2, '合同 (Contract)', [
        'contract_id (PK)', 'file_url',
        'sign_date', 'content'
    ], pk_idx=0, width=2.2, color='#E0F2F1', edge='#00695C')

    # Relationship: 引导申报 (Notice -> Project, 1:n)
    draw_relationship(ax, 3.5, 2.2, '引导\n申报', width=0.9, height=0.55)
    ax.annotate('', xy=(4.25, 2.2), xytext=(4.05, 2.2),
                arrowprops=dict(arrowstyle='-', color='#37474F', lw=1.2))
    ax.annotate('', xy=(2.8, 4.2), xytext=(3.05, 2.4),
                arrowprops=dict(arrowstyle='-', color='#37474F', lw=1.2))
    draw_cardinality(ax, 2.4, 3.3, '1')
    draw_cardinality(ax, 4.3, 2.5, 'n')

    # Relationship: 评审 (Researcher <-> Project, m:n)
    draw_relationship(ax, 5.5, 3.7, '评审', width=0.9, height=0.5)
    ax.annotate('', xy=(5.5, 3.95), xytext=(5.5, 4.5),
                arrowprops=dict(arrowstyle='-', color='#37474F', lw=1.2))
    ax.annotate('', xy=(5.5, 2.95), xytext=(5.5, 3.45),
                arrowprops=dict(arrowstyle='-', color='#37474F', lw=1.2))
    draw_cardinality(ax, 5.15, 4.7, 'm')
    draw_cardinality(ax, 5.15, 2.5, 'n')
    # Review attributes
    ax.text(3.8, 3.7, '评审结果\n打分(0-100)\n评语', ha='center', va='center',
            fontsize=6.5, color='#6A1B9A', style='italic')
    ax.annotate('', xy=(5.05, 3.7), xytext=(4.35, 3.7),
                arrowprops=dict(arrowstyle='-', color='#6A1B9A', lw=0.8))

    # Relationship: 签署 (Project <-> Contract, 1:1)
    draw_relationship(ax, 7.3, 2.2, '签署', width=0.75, height=0.45)
    ax.annotate('', xy=(7.7, 2.2), xytext=(6.8, 2.2),
                arrowprops=dict(arrowstyle='-', color='#37474F', lw=1.2))
    ax.annotate('', xy=(8.0, 2.2), xytext=(7.7, 2.2),
                arrowprops=dict(arrowstyle='-', color='#37474F', lw=1.2))
    draw_cardinality(ax, 6.9, 2.5, '1')
    draw_cardinality(ax, 7.9, 2.5, '1')

    plt.tight_layout(pad=0.5)
    save_fig(fig, 'Paper/Picture/fig3_2_局部ER2.png')
    add_figure('Paper/Picture/fig3_2_局部ER2.png', '图3-2  局部E-R图二：项目审批与评审管理')
except Exception as e:
    print(f"WARN fig3-2: {e}")
    body('（图3-2：局部E-R图二——待补充）')


# ---- 3.1.3 局部E-R图3: 经费与成果管理 ----
doc.add_heading('3.1.3  局部E-R图三：经费与成果管理', level=3)

body(
    '该局部E-R图覆盖经费管理与成果管理两大业务域。科研项目（Project）与预算科目（Budget）'
    '之间为1:n的"编制"联系（一个项目编制多个预算科目，每个预算科目仅属于一个项目）。'
    '预算科目与经费流水（Expenditure）之间为1:n的"记录"联系（一个预算科目可有多条经费'
    '流水记录）。经费流水与研究人员之间存在两个联系：财务处管理员执行"登记/审批"操作'
    '（1:n），科研人员发起"报销申请"操作（1:n）。科研项目与科研成果（Achievement）之间'
    '为1:n的"产出"联系（一个项目可产出多项成果，一项成果仅属于一个项目）。'
)

try:
    fig, ax = plt.subplots(figsize=(11, 7))
    ax.set_xlim(0, 11); ax.set_ylim(0, 7); ax.axis('off')
    ax.set_facecolor('white'); fig.patch.set_facecolor('white')
    ax.text(5.5, 6.7, '局部E-R图三：经费与成果管理', ha='center', va='center',
            fontsize=12, fontweight='bold', color='#1a1a2e')

    # Entity: Project
    draw_entity(ax, 1.5, 4.5, '科研项目 (Project)', [
        'project_id (PK)', 'name', 'budget_total', 'status'
    ], pk_idx=0, width=2.2, color='#FFF3E0', edge='#E65100')

    # Entity: Budget
    draw_entity(ax, 1.5, 1.8, '预算科目 (Budget)', [
        'budget_id (PK)', 'category', 'amount', 'spent'
    ], pk_idx=0, width=2.2, color='#E3F2FD', edge='#1565C0')

    # Entity: Expenditure
    draw_entity(ax, 5.5, 3.2, '经费流水 (Expenditure)', [
        'exp_id (PK)', 'type', 'amount', 'exp_date',
        'purpose', 'approval_status'
    ], pk_idx=0, width=2.4, color='#FCE4EC', edge='#C62828')

    # Entity: Achievement
    draw_entity(ax, 5.5, 5.8, '科研成果 (Achievement)', [
        'ach_id (PK)', 'type', 'title',
        'publish_date', 'author'
    ], pk_idx=0, width=2.4, color='#E8F5E9', edge='#2E7D32')

    # Entity: Researcher (simplified)
    draw_entity(ax, 9.2, 3.2, '研究人员\n(Researcher)', [
        'researcher_id (PK)', 'name', 'role'
    ], pk_idx=0, width=2.0, color='#ECEFF1', edge='#546E7A')

    # Relationship: 编制 (Project -> Budget, 1:n)
    draw_relationship(ax, 1.5, 3.15, '编制', width=0.8, height=0.5)
    ax.annotate('', xy=(1.5, 3.4), xytext=(1.5, 3.75),
                arrowprops=dict(arrowstyle='-', color='#37474F', lw=1.2))
    ax.annotate('', xy=(1.5, 2.55), xytext=(1.5, 2.9),
                arrowprops=dict(arrowstyle='-', color='#37474F', lw=1.2))
    draw_cardinality(ax, 1.8, 4.0, '1')
    draw_cardinality(ax, 1.8, 2.3, 'n')

    # Relationship: 记录 (Budget -> Expenditure, 1:n)
    draw_relationship(ax, 3.5, 1.8, '记录', width=0.75, height=0.45)
    ax.annotate('', xy=(4.3, 2.4), xytext=(3.85, 1.95),
                arrowprops=dict(arrowstyle='-', color='#37474F', lw=1.2))
    ax.annotate('', xy=(3.1, 1.95), xytext=(3.85, 1.95),
                arrowprops=dict(arrowstyle='-', color='#37474F', lw=1.2))
    draw_cardinality(ax, 3.3, 1.5, '1')
    draw_cardinality(ax, 4.0, 2.6, 'n')

    # Relationship: 产出 (Project -> Achievement, 1:n)
    draw_relationship(ax, 3.5, 5.8, '产出', width=0.75, height=0.45)
    ax.annotate('', xy=(3.9, 5.8), xytext=(2.65, 5.8),
                arrowprops=dict(arrowstyle='-', color='#37474F', lw=1.2))
    ax.annotate('', xy=(4.3, 5.8), xytext=(3.9, 5.8),
                arrowprops=dict(arrowstyle='-', color='#37474F', lw=1.2))
    draw_cardinality(ax, 2.7, 5.5, '1')
    draw_cardinality(ax, 4.2, 5.5, 'n')

    # Relationship: 登记/审批 (Researcher -> Expenditure, 1:n)
    draw_relationship(ax, 7.5, 3.2, '登记/\n审批', width=0.8, height=0.55)
    ax.annotate('', xy=(7.9, 3.2), xytext=(6.75, 3.2),
                arrowprops=dict(arrowstyle='-', color='#37474F', lw=1.2))
    ax.annotate('', xy=(8.2, 3.2), xytext=(7.9, 3.2),
                arrowprops=dict(arrowstyle='-', color='#37474F', lw=1.2))
    draw_cardinality(ax, 6.9, 3.5, 'n')
    draw_cardinality(ax, 8.5, 3.5, '1')

    # Relationship: 报销申请 (Researcher -> Expenditure, 1:n)
    draw_relationship(ax, 7.5, 1.6, '报销\n申请', width=0.8, height=0.55)
    ax.annotate('', xy=(6.75, 2.45), xytext=(7.1, 1.8),
                arrowprops=dict(arrowstyle='-', color='#37474F', lw=1.2))
    ax.annotate('', xy=(8.2, 1.6), xytext=(7.9, 1.6),
                arrowprops=dict(arrowstyle='-', color='#37474F', lw=1.2))
    draw_cardinality(ax, 7.0, 1.35, 'n')
    draw_cardinality(ax, 8.5, 1.35, '1')

    plt.tight_layout(pad=0.5)
    save_fig(fig, 'Paper/Picture/fig3_3_局部ER3.png')
    add_figure('Paper/Picture/fig3_3_局部ER3.png', '图3-3  局部E-R图三：经费与成果管理')
except Exception as e:
    print(f"WARN fig3-3: {e}")
    body('（图3-3：局部E-R图三——待补充）')


# ===================== 3.2 全局E-R图设计 =====================
doc.add_heading('3.2  全局E-R图设计', level=2)

body(
    '全局E-R图是将所有局部E-R图合并为一个完整的概念模型，展示系统中全部实体及其联系。'
    '合并过程需要识别并消解三类冲突：属性冲突、命名冲突和结构冲突。以下先阐述冲突消解'
    '过程，再呈现最终的全局E-R图。'
)

# 3.2.1 冲突消解
doc.add_heading('3.2.1  冲突消解', level=3)

body(
    '（1）属性冲突。属性冲突包括属性域冲突（同一属性在不同局部E-R图中取值范围不一致）'
    '和属性取值单位冲突。例如，"项目类型"在局部E-R图一中使用ENUM类型（纵向/横向），在'
    '局部E-R图二中可能被定义为VARCHAR，需要统一为ENUM类型并约定取值集合为{纵向, 横向}。'
    '"预算金额"统一以"万元"为计量单位，采用DECIMAL(12,2)数据类型，避免不同局部视图'
    '中的单位不一致。'
)

body(
    '（2）命名冲突。命名冲突包括同名异义（同一名称在不同局部E-R图中表示不同对象）和'
    '异名同义（不同名称表示同一对象）。例如，"研究人员"实体在局部E-R图一中全称为'
    '"研究人员（Researcher）"，在局部E-R图二的评审场景中可能被称为"评审专家"，但实际'
    '上评审专家是研究人员实体的一个子集（role=专家）。全局E-R图中统一为"研究人员"实体，'
    '通过role属性区分角色。再如，"项目编号"在不同局部视图中可能称为project_id或pid，'
    '全局统一为project_id。'
)

body(
    '（3）结构冲突。结构冲突包括同一对象在不同局部E-R图中被抽象为不同层次——在一个视图中'
    '作为实体，在另一个视图中作为属性；或者同一联系在不同视图中具有不同的联系类型。'
    '在本系统中，经费流水（Expenditure）在局部E-R图三中作为独立实体存在，因其具有自身'
    '的属性集合（金额、日期、审批状态等）且参与多个联系。全局E-R图中保留其独立实体地位。'
    '此外，研究人员与科研项目之间的"参与"联系在局部E-R图一中建模为m:n联系，全局视图中'
    '维持该设定。所有冲突消解完毕后，得到最终全局E-R图如下。'
)

# 3.2.2 全局E-R图
doc.add_heading('3.2.2  全局E-R图', level=3)

body(
    '全局E-R图集成了三个局部视图的全部实体与联系，共包含11类实体、14个联系，完整刻画了'
    '高校科研项目管理系统的信息模型。全局E-R图是下一章逻辑结构设计的直接输入。'
)

try:
    fig, ax = plt.subplots(figsize=(13, 8.5))
    ax.set_xlim(0, 13); ax.set_ylim(0, 8.5); ax.axis('off')
    ax.set_facecolor('white'); fig.patch.set_facecolor('white')
    ax.text(6.5, 8.2, '全局E-R图：高校科研项目管理系统', ha='center', va='center',
            fontsize=13, fontweight='bold', color='#1a1a2e')

    # ===== ENTITIES (positioned to avoid crossing lines) =====
    # Top row
    draw_entity(ax, 2.0, 6.8, '学院\n(College)', [
        'college_id (PK)', 'name', 'dean', 'tel'
    ], pk_idx=0, width=2.0, color='#E3F2FD', edge='#1565C0')

    draw_entity(ax, 5.5, 6.8, '通知公告\n(Notice)', [
        'notice_id (PK)', 'title', 'content',
        'publish_date'
    ], pk_idx=0, width=2.0, color='#F3E5F5', edge='#6A1B9A')

    draw_entity(ax, 9.0, 6.8, '合同\n(Contract)', [
        'contract_id (PK)', 'file_url',
        'sign_date', 'content'
    ], pk_idx=0, width=2.0, color='#E0F2F1', edge='#00695C')

    # Middle row
    draw_entity(ax, 2.0, 4.2, '研究人员\n(Researcher)', [
        'researcher_id (PK)', 'name', 'password',
        'title', 'phone', 'email', 'role'
    ], pk_idx=0, width=2.2, color='#E8F5E9', edge='#2E7D32')

    draw_entity(ax, 6.5, 4.2, '科研项目\n(Project)', [
        'project_id (PK)', 'name', 'type',
        'level', 'start_date', 'end_date',
        'budget_total', 'status'
    ], pk_idx=0, width=2.3, color='#FFF3E0', edge='#E65100')

    draw_entity(ax, 10.5, 4.2, '评审记录\n(Review)', [
        '(project_id, expert_id) (PK)',
        'result', 'score', 'comment'
    ], pk_idx=0, width=2.4, color='#FFF9C4', edge='#F57F17')

    # Bottom row
    draw_entity(ax, 2.0, 1.5, '经费流水\n(Expenditure)', [
        'exp_id (PK)', 'type', 'amount',
        'exp_date', 'purpose', 'approval_status'
    ], pk_idx=0, width=2.2, color='#FCE4EC', edge='#C62828')

    draw_entity(ax, 5.5, 1.5, '预算科目\n(Budget)', [
        'budget_id (PK)', 'category',
        'amount', 'spent'
    ], pk_idx=0, width=2.0, color='#E3F2FD', edge='#1565C0')

    draw_entity(ax, 8.5, 1.5, '科研成果\n(Achievement)', [
        'ach_id (PK)', 'type', 'title',
        'publish_date', 'author'
    ], pk_idx=0, width=2.2, color='#E8F5E9', edge='#2E7D32')

    draw_entity(ax, 11.2, 1.5, '进展报告\n(ProgressReport)', [
        'report_id (PK)', 'year',
        'submit_date', 'content'
    ], pk_idx=0, width=2.2, color='#ECEFF1', edge='#546E7A')

    # ===== RELATIONSHIPS =====
    # College --归属--> Researcher (1:n)
    draw_relationship(ax, 3.0, 5.5, '归属', width=0.7, height=0.4)
    ax.plot([2.7, 3.0], [6.0, 5.7], color='#37474F', lw=1.0)
    ax.plot([2.7, 3.0], [5.0, 5.3], color='#37474F', lw=1.0)
    draw_cardinality(ax, 2.5, 5.9, '1')
    draw_cardinality(ax, 2.5, 4.9, 'n')

    # Researcher --负责--> Project (1:n)
    draw_relationship(ax, 4.2, 4.8, '负责', width=0.7, height=0.4)
    ax.plot([3.2, 3.9], [4.5, 4.7], color='#37474F', lw=1.0)
    ax.plot([5.3, 3.9], [4.5, 4.7], color='#37474F', lw=1.0)
    draw_cardinality(ax, 3.4, 4.3, '1')
    draw_cardinality(ax, 5.2, 4.3, 'n')

    # Researcher --参与--> Project (m:n)
    draw_relationship(ax, 4.2, 3.5, '参与', width=0.7, height=0.4)
    ax.plot([3.2, 3.9], [3.9, 3.7], color='#37474F', lw=1.0)
    ax.plot([5.3, 3.9], [3.9, 3.7], color='#37474F', lw=1.0)
    draw_cardinality(ax, 3.5, 4.0, 'm')
    draw_cardinality(ax, 5.0, 4.0, 'n')

    # Researcher --评审--> Project (m:n, via Review)
    # Review is the associative entity for m:n relationship
    ax.plot([7.8, 5.7], [4.2, 4.2], color='#37474F', lw=1.0)
    ax.plot([9.3, 6.5], [4.2, 4.5], color='#37474F', lw=1.0)
    draw_cardinality(ax, 7.0, 4.5, 'm')
    draw_cardinality(ax, 7.0, 3.9, 'n')

    # Notice --引导申报--> Project (1:n)
    draw_relationship(ax, 6.5, 5.8, '引导\n申报', width=0.75, height=0.45)
    ax.plot([6.0, 6.2], [6.3, 6.0], color='#37474F', lw=1.0)
    ax.plot([6.0, 6.2], [5.0, 5.6], color='#37474F', lw=1.0)
    draw_cardinality(ax, 5.7, 6.2, '1')
    draw_cardinality(ax, 5.7, 5.0, 'n')

    # Project --签署--> Contract (1:1)
    draw_relationship(ax, 8.0, 5.8, '签署', width=0.65, height=0.38)
    ax.plot([7.7, 7.8], [5.0, 5.6], color='#37474F', lw=1.0)
    ax.plot([8.5, 7.8], [6.3, 5.9], color='#37474F', lw=1.0)
    draw_cardinality(ax, 7.5, 5.0, '1')
    draw_cardinality(ax, 8.5, 6.2, '1')

    # Project --编制--> Budget (1:n)
    draw_relationship(ax, 6.5, 2.8, '编制', width=0.7, height=0.4)
    ax.plot([6.0, 6.2], [3.6, 3.0], color='#37474F', lw=1.0)
    ax.plot([6.0, 6.2], [2.1, 2.6], color='#37474F', lw=1.0)
    draw_cardinality(ax, 5.7, 3.5, '1')
    draw_cardinality(ax, 5.7, 2.0, 'n')

    # Budget --记录--> Expenditure (1:n)
    draw_relationship(ax, 3.8, 2.0, '记录', width=0.65, height=0.38)
    ax.plot([3.1, 3.5], [2.0, 2.0], color='#37474F', lw=1.0)
    ax.plot([3.1, 3.5], [1.8, 1.8], color='#37474F', lw=1.0)
    draw_cardinality(ax, 2.8, 1.7, '1')
    draw_cardinality(ax, 3.2, 1.3, 'n')

    # Researcher --报销申请--> Expenditure (1:n)
    ax.plot([3.2, 2.8], [3.8, 2.5], color='#37474F', lw=1.0)
    draw_cardinality(ax, 3.5, 3.2, '1')
    draw_cardinality(ax, 2.5, 2.5, 'n')

    # Researcher --登记/审批--> Expenditure (1:n)
    ax.plot([2.5, 2.5], [3.8, 2.5], color='#37474F', lw=1.0)
    draw_cardinality(ax, 2.2, 3.2, '1')
    draw_cardinality(ax, 2.2, 2.5, 'n')

    # Project --产出--> Achievement (1:n)
    draw_relationship(ax, 8.5, 2.8, '产出', width=0.65, height=0.38)
    ax.plot([7.7, 8.2], [3.6, 3.0], color='#37474F', lw=1.0)
    ax.plot([8.0, 8.2], [2.1, 2.6], color='#37474F', lw=1.0)
    draw_cardinality(ax, 7.5, 3.5, '1')
    draw_cardinality(ax, 7.7, 2.0, 'n')

    # Project --提交--> ProgressReport (1:n)
    draw_relationship(ax, 9.8, 2.8, '提交', width=0.65, height=0.38)
    ax.plot([8.3, 9.5], [3.6, 3.0], color='#37474F', lw=1.0)
    ax.plot([10.5, 9.8], [2.1, 2.6], color='#37474F', lw=1.0)
    draw_cardinality(ax, 8.2, 3.5, '1')
    draw_cardinality(ax, 10.5, 2.0, 'n')

    # Researcher --提交报告--> ProgressReport (1:n)
    ax.plot([3.2, 9.8], [3.8, 2.5], color='#37474F', lw=1.0)
    draw_cardinality(ax, 4.5, 3.4, '1')
    draw_cardinality(ax, 9.0, 2.7, 'n')

    plt.tight_layout(pad=0.3)
    save_fig(fig, 'Paper/Picture/fig3_4_全局ER.png')
    add_figure('Paper/Picture/fig3_4_全局ER.png', '图3-4  全局E-R图', width=6.2)
except Exception as e:
    print(f"WARN fig3-4: {e}")
    import traceback; traceback.print_exc()
    body('（图3-4：全局E-R图——待补充）')


# ===================== 3.3 实体与联系汇总 =====================
doc.add_heading('3.3  实体与联系汇总', level=2)

body(
    '为便于后续逻辑结构设计，将全局E-R图中的全部实体与联系进行汇总。表3-1列出了系统中'
    '的11类实体及其核心属性，表3-2列出了14个联系及其类型和参与实体。'
)

doc.add_heading('3.3.1  实体清单', level=3)

make_table(
    ['序号', '实体名称', '对应表名', '主键', '属性数量', '说明'],
    [
        ['E01', '学院', 'college', 'college_id', '4', '大学的二级学院/系所'],
        ['E02', '研究人员', 'researcher', 'researcher_id', '7', '含科研人员/专家/科研处/财务处'],
        ['E03', '科研项目', 'project', 'project_id', '8', '核心实体，10种状态流转'],
        ['E04', '预算科目', 'budget', 'budget_id', '4', '项目的经费预算明细'],
        ['E05', '经费流水', 'expenditure', 'exp_id', '6', '每笔到账或支出的记录'],
        ['E06', '评审记录', 'review', '(project_id,expert_id)', '4', 'm:n联系的关联实体'],
        ['E07', '通知公告', 'notice', 'notice_id', '4', '科研处发布的申报通知'],
        ['E08', '合同', 'contract', 'contract_id', '4', '项目合同/任务书'],
        ['E09', '进展报告', 'progress_report', 'report_id', '4', '按年度提交的进展报告'],
        ['E10', '验收记录', 'acceptance', 'acceptance_id', '5', '验收申请与评审结果'],
        ['E11', '科研成果', 'achievement', 'ach_id', '5', '6种类型：论文/专利/软著等'],
    ]
)

doc.add_heading('3.3.2  联系清单', level=3)

make_table(
    ['序号', '联系名称', '类型', '参与实体', '联系的属性', '说明'],
    [
        ['R01', '归属', '1:n', '学院→研究人员', '—', '一个学院有多名研究人员'],
        ['R02', '负责', '1:n', '研究人员→科研项目', '—', '一个项目有一名负责人'],
        ['R03', '参与', 'm:n', '研究人员↔科研项目', '参与角色', '多人参与一个项目'],
        ['R04', '评审', 'm:n', '研究人员↔科研项目', '结果/打分/评语', '专家评审项目(关联实体Review)'],
        ['R05', '发布', '1:n', '研究人员→通知公告', '—', '科研处发布通知'],
        ['R06', '引导申报', '1:n', '通知公告→科研项目', '—', '通知引导项目申报'],
        ['R07', '签署', '1:1', '科研项目↔合同', '—', '一个项目签一份合同'],
        ['R08', '编制', '1:n', '科研项目→预算科目', '—', '项目编制多个预算科目'],
        ['R09', '记录', '1:n', '预算科目→经费流水', '—', '科目下的经费流水'],
        ['R10', '登记/审批', '1:n', '研究人员→经费流水', '—', '财务处登记审批经费'],
        ['R11', '报销申请', '1:n', '研究人员→经费流水', '—', '科研人员提交报销'],
        ['R12', '产出', '1:n', '科研项目→科研成果', '—', '项目产出多项成果'],
        ['R13', '提交', '1:n', '科研项目→进展报告', '—', '项目按年度提交报告'],
        ['R14', '验收', '1:1', '科研项目↔验收记录', '—', '一个项目一条验收记录'],
    ]
)


# ---- Save ----
output_path = 'Paper/第3章_概念结构设计_v1.docx'
doc.save(output_path)
print(f"\n{'='*60}")
print("[DONE] Chapter 3 v1 generated successfully")
print(f"File: {output_path}")
print(f"{'='*60}")
