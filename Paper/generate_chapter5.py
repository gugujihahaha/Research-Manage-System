# -*- coding: utf-8 -*-
"""生成第5章 数据库物理结构设计 DOCX — 精修版
包含: 12张表结构 / 字段约束 / 存储引擎与字符集
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, sys

sys.stdout.reconfigure(encoding='utf-8', errors='replace')

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54); section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17); section.right_margin = Cm(3.17)

style = doc.styles['Normal']
style.font.name = '宋体'; style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Cm(0.74)

for lvl, sz in [(1, 16), (2, 14), (3, 13)]:
    s = doc.styles[f'Heading {lvl}']
    s.font.name = '黑体'; s.font.size = Pt(sz); s.font.bold = True
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if lvl == 1 else WD_ALIGN_PARAGRAPH.LEFT
    s.paragraph_format.space_before = Pt(12 if lvl == 1 else 8)
    s.paragraph_format.space_after = Pt(6 if lvl == 1 else 4)


def body(text):
    p = doc.add_paragraph(style='Normal'); p.clear()
    r = p.add_run(text)
    r.font.name = '宋体'; r.font.size = Pt(12)
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    return p


def make_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.font.name = '宋体'; r.font.size = Pt(9); r.font.bold = True
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), '4472C4'); shd.set(qn('w:val'), 'clear')
        c._tc.get_or_add_tcPr().append(shd)
        r.font.color.rgb = RGBColor(255, 255, 255)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.rows[i + 1].cells[j]; c.text = ''
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val)); r.font.name = '宋体'; r.font.size = Pt(8)
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if i % 2 == 0:
                shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'F2F6FC'); shd.set(qn('w:val'), 'clear')
                c._tc.get_or_add_tcPr().append(shd)
    doc.add_paragraph()
    return t


def table_schema(title, desc, fields):
    """Create a standard table schema definition.
    fields: list of (field_name, type_len, nullable, default, constraints, description)
    """
    doc.add_heading(title, level=3)
    body(desc)
    make_table(
        ['字段名', '数据类型/长度', '允许为空', '默认值', '约束', '说明'],
        fields
    )


# ================================================================
#  CHAPTER 5
# ================================================================
doc.add_heading('第5章  数据库物理结构设计', level=1)

body(
    '物理结构设计是在逻辑结构设计的基础上，为每个关系模式确定其在具体数据库管理系统'
    '（MySQL 8.0）中的物理存储方案。本章的核心任务包括：逐一确定每张数据表的字段名、'
    '数据类型与长度、是否允许空值、默认值配置及各类完整性约束（PRIMARY KEY、FOREIGN KEY、'
    'CHECK、UNIQUE、DEFAULT）；说明存储引擎（InnoDB）与字符集（utf8mb4）的选择依据；'
    '定义索引的物理存储方式。物理结构设计直接产生可在MySQL中执行的DDL建表语句，是第6章'
    '数据库实施的直接前置工作。'
)

# ===================== 5.1 表结构设计 =====================
doc.add_heading('5.1  表结构设计', level=2)

body(
    '本节依据第4章逻辑结构设计中确定的12个关系模式，逐一设计每张数据表的物理存储结构。'
    '数据类型选择遵循以下原则：（1）字符型字段选用VARCHAR以节约存储空间，仅在定长编码'
    '字段（如college_id、researcher_id）使用CHAR；（2）金额字段统一使用DECIMAL(12,2)'
    '以保证精度，避免浮点误差；（3）日期字段使用DATE类型，仅记录日期信息无需时间部分；'
    '（4）枚举值字段（如type、status、role）使用ENUM类型，在数据库层面限制非法取值；'
    '（5）自增主键使用INT AUTO_INCREMENT。'
)

# ---- 5.1.1 college ----
table_schema(
    '5.1.1  学院表（college）',
    '存储各学院的机构信息，是研究人员表的上级参照表。',
    [
        ('college_id', 'CHAR(4)', '否', '—', 'PRIMARY KEY', '学院编号，C001~C006'),
        ('name', 'VARCHAR(50)', '否', '—', '—', '学院全称'),
        ('dean', 'VARCHAR(20)', '是', 'NULL', '—', '院长姓名'),
        ('tel', 'VARCHAR(20)', '是', 'NULL', '—', '学院联系电话'),
    ]
)

# ---- 5.1.2 researcher ----
table_schema(
    '5.1.2  研究人员表（researcher）',
    '存储系统所有用户信息，角色字段区分科研人员、评审专家、科研处管理员和财务处管理员。'
    'college_id外键关联学院表，删除学院时设为NULL以保留用户记录。',
    [
        ('researcher_id', 'CHAR(10)', '否', '—', 'PRIMARY KEY', '工号，R开头+4位数字，如R0001'),
        ('name', 'VARCHAR(20)', '否', '—', '—', '用户姓名'),
        ('password', 'VARCHAR(100)', '否', '123456', 'DEFAULT', '登录密码，当前为明文（待改为bcrypt加密）'),
        ('title', 'VARCHAR(20)', '是', 'NULL', '—', '职称：教授/副教授/讲师/助教'),
        ('college_id', 'CHAR(4)', '是', 'NULL', 'FOREIGN KEY→college(college_id) ON DELETE SET NULL', '所属学院编号'),
        ('phone', 'VARCHAR(20)', '是', 'NULL', '—', '联系电话'),
        ('email', 'VARCHAR(50)', '是', 'NULL', '—', '电子邮箱'),
        ('role', "ENUM('科研人员','专家','科研处','财务处')", '否', '科研人员', 'DEFAULT', '用户角色'),
    ]
)

# ---- 5.1.3 project ----
table_schema(
    '5.1.3  科研项目表（project）',
    '系统的核心数据表，存储所有科研项目信息。status字段定义了10种项目状态，覆盖从申报到'
    '验收的全生命周期。leader_id外键关联研究人员表。',
    [
        ('project_id', 'CHAR(20)', '否', '—', 'PRIMARY KEY', '项目编号，格式P+年份+序号，如P2024001'),
        ('name', 'VARCHAR(200)', '否', '—', '—', '项目名称'),
        ('type', "ENUM('纵向','横向')", '否', '—', '—', '项目类型'),
        ('level', 'VARCHAR(20)', '是', 'NULL', '—', '级别：国家级/省部级/市厅级/横向'),
        ('leader_id', 'CHAR(10)', '否', '—', 'FOREIGN KEY→researcher(researcher_id)', '项目负责人工号'),
        ('apply_date', 'DATE', '是', 'NULL', '—', '申请日期'),
        ('start_date', 'DATE', '是', 'NULL', '—', '项目开始日期'),
        ('end_date', 'DATE', '是', 'NULL', '—', '项目结束日期'),
        ('budget_total', 'DECIMAL(12,2)', '是', 'NULL', '—', '总预算（万元）'),
        ('status', "ENUM('申报中','形式审查','专家评审','立项公示','已立项','执行中','验收申请','验收评审','验收通过','终止')",
         '否', '申报中', 'DEFAULT', '项目当前状态，10种状态流转'),
        ('acceptance_date', 'DATE', '是', 'NULL', '—', '验收日期，由触发器自动填充'),
        ('file_url', 'VARCHAR(255)', '是', 'NULL', '—', '附件文件路径'),
    ]
)

# ---- 5.1.4 budget ----
table_schema(
    '5.1.4  预算科目表（budget）',
    '存储每个项目的预算科目明细，一个项目可编制多个预算科目。project_id外键关联项目表，'
    '删除项目时级联删除预算科目。',
    [
        ('budget_id', 'INT', '否', '—', 'PRIMARY KEY AUTO_INCREMENT', '科目ID，自增'),
        ('project_id', 'CHAR(20)', '否', '—', 'FOREIGN KEY→project(project_id) ON DELETE CASCADE', '所属项目编号'),
        ('category', 'VARCHAR(50)', '否', '—', '—', '科目名称：设备费/材料费/差旅费/劳务费/会议费/其他'),
        ('amount', 'DECIMAL(12,2)', '否', '—', '—', '预算金额（万元）'),
        ('spent', 'DECIMAL(12,2)', '否', '0.00', 'DEFAULT 0', '已支出金额（万元），由触发器自动更新'),
    ]
)

# ---- 5.1.5 expenditure ----
table_schema(
    '5.1.5  经费流水表（expenditure）',
    '记录每一笔科研经费的到账或支出明细。type字段区分"到账"与"支出"，approval_status字段'
    '跟踪审批流程。三个外键分别关联项目表、预算科目表和研究人员表。',
    [
        ('exp_id', 'INT', '否', '—', 'PRIMARY KEY AUTO_INCREMENT', '流水ID，自增'),
        ('project_id', 'CHAR(20)', '否', '—', 'FOREIGN KEY→project(project_id)', '所属项目编号'),
        ('budget_id', 'INT', '是', 'NULL', 'FOREIGN KEY→budget(budget_id)', '关联预算科目（支出时必填）'),
        ('type', "ENUM('到账','支出')", '否', '—', '—', '交易类型'),
        ('amount', 'DECIMAL(12,2)', '否', '—', '—', '金额（万元）'),
        ('exp_date', 'DATE', '否', '—', '—', '交易日期'),
        ('purpose', 'TEXT', '是', 'NULL', '—', '用途说明'),
        ('operator_id', 'CHAR(10)', '是', 'NULL', 'FOREIGN KEY→researcher(researcher_id)', '经办人/操作员工号'),
        ('approval_status', "ENUM('待审批','已通过','已驳回')", '否', '待审批', 'DEFAULT', '审批状态'),
    ]
)

# ---- 5.1.6 review ----
table_schema(
    '5.1.6  评审记录表（review）',
    '存储专家对项目的评审意见，由m:n联系转换而来。复合主键(project_id, expert_id)确保'
    '每位专家对一个项目仅有一条评审记录。result字段取值约束了评审结果的合法性。',
    [
        ('project_id', 'CHAR(20)', '否', '—', 'PRIMARY KEY (复合)', '被评审项目编号'),
        ('expert_id', 'CHAR(10)', '否', '—', 'PRIMARY KEY (复合)', '评审专家工号'),
        ('result', "ENUM('通过','不通过')", '是', 'NULL', '—', '评审结果'),
        ('score', 'INT', '是', 'NULL', 'CHECK(score>=0 AND score<=100)', '评审打分，0-100分'),
        ('comment', 'TEXT', '是', 'NULL', '—', '评审评语'),
        ('review_date', 'DATE', '是', 'NULL', '—', '评审日期'),
    ]
)

# ---- 5.1.7 notice ----
table_schema(
    '5.1.7  通知公告表（notice）',
    '存储科研处发布的项目申报通知。按发布时间倒序查询，支持所有用户查看。',
    [
        ('notice_id', 'INT', '否', '—', 'PRIMARY KEY AUTO_INCREMENT', '通知ID，自增'),
        ('title', 'VARCHAR(200)', '否', '—', '—', '通知标题'),
        ('content', 'TEXT', '否', '—', '—', '通知正文内容'),
        ('publish_date', 'DATE', '否', '—', '—', '发布日期'),
        ('publisher_id', 'CHAR(10)', '是', 'NULL', 'FOREIGN KEY→researcher(researcher_id)', '发布人工号'),
    ]
)

# ---- 5.1.8 contract ----
table_schema(
    '5.1.8  合同表（contract）',
    '存储项目合同/任务书信息，与项目一对一关联。project_id设置为UNIQUE约束确保一一对应。',
    [
        ('contract_id', 'INT', '否', '—', 'PRIMARY KEY AUTO_INCREMENT', '合同ID，自增'),
        ('project_id', 'CHAR(20)', '否', '—', 'FOREIGN KEY→project(project_id) UNIQUE', '关联项目编号，唯一约束保证1:1'),
        ('file_url', 'VARCHAR(255)', '是', 'NULL', '—', '合同文件存储路径'),
        ('sign_date', 'DATE', '是', 'NULL', '—', '合同签署日期'),
        ('content', 'TEXT', '是', 'NULL', '—', '合同主要内容描述'),
    ]
)

# ---- 5.1.9 progress_report ----
table_schema(
    '5.1.9  进展报告表（progress_report）',
    '存储科研人员按年度提交的项目进展报告，初始状态为"待审核"由科研处审核。',
    [
        ('report_id', 'INT', '否', '—', 'PRIMARY KEY AUTO_INCREMENT', '报告ID，自增'),
        ('project_id', 'CHAR(20)', '否', '—', 'FOREIGN KEY→project(project_id)', '关联项目编号'),
        ('report_year', 'INT', '否', '—', '—', '报告年度，如2024'),
        ('submit_date', 'DATE', '否', '—', '—', '提交日期'),
        ('content', 'TEXT', '否', '—', '—', '报告正文内容'),
        ('review_status', "ENUM('待审核','已通过','已驳回')", '否', '待审核', 'DEFAULT', '审核状态'),
    ]
)

# ---- 5.1.10 change_request ----
table_schema(
    '5.1.10  变更申请表（change_request）',
    '存储项目执行过程中的变更申请，支持三种变更类型。old_value和new_value分别记录变更前后的值。',
    [
        ('change_id', 'INT', '否', '—', 'PRIMARY KEY AUTO_INCREMENT', '变更ID，自增'),
        ('project_id', 'CHAR(20)', '否', '—', 'FOREIGN KEY→project(project_id)', '关联项目编号'),
        ('change_type', "ENUM('预算调整','延期','成员变更')", '否', '—', '—', '变更类型'),
        ('old_value', 'VARCHAR(200)', '是', 'NULL', '—', '变更前的值'),
        ('new_value', 'VARCHAR(200)', '否', '—', '—', '变更后的值'),
        ('reason', 'TEXT', '否', '—', '—', '变更理由说明'),
        ('request_date', 'DATE', '否', '—', '—', '申请日期'),
        ('approval_status', "ENUM('待审批','已通过','已驳回')", '否', '待审批', 'DEFAULT', '审批状态'),
    ]
)

# ---- 5.1.11 acceptance ----
table_schema(
    '5.1.11  验收记录表（acceptance）',
    '存储项目验收申请与评审结果信息。acceptance_date由触发器trg_acceptance_update在'
    '评审通过时自动填充。',
    [
        ('acceptance_id', 'INT', '否', '—', 'PRIMARY KEY AUTO_INCREMENT', '验收ID，自增'),
        ('project_id', 'CHAR(20)', '否', '—', 'FOREIGN KEY→project(project_id)', '关联项目编号'),
        ('apply_date', 'DATE', '否', '—', '—', '验收申请提交日期'),
        ('review_result', "ENUM('通过','不通过')", '是', 'NULL', '—', '验收评审结果'),
        ('acceptance_date', 'DATE', '是', 'NULL', '—', '验收通过日期，触发器自动填充'),
        ('certificate_url', 'VARCHAR(255)', '是', 'NULL', '—', '结题证书文件链接'),
    ]
)

# ---- 5.1.12 achievement ----
table_schema(
    '5.1.12  科研成果表（achievement）',
    '存储科研人员登记的各类科研成果，支持六种成果类型。review_status字段支持科研处审核流程。',
    [
        ('ach_id', 'INT', '否', '—', 'PRIMARY KEY AUTO_INCREMENT', '成果ID，自增'),
        ('project_id', 'CHAR(20)', '否', '—', 'FOREIGN KEY→project(project_id)', '关联项目编号'),
        ('type', "ENUM('论文','专利','软件著作权','获奖','标准','成果转化')", '否', '—', '—', '成果类型'),
        ('title', 'VARCHAR(200)', '否', '—', '—', '成果标题'),
        ('publish_date', 'DATE', '是', 'NULL', '—', '发表日期'),
        ('author', 'VARCHAR(100)', '是', 'NULL', '—', '作者/完成人'),
        ('file_url', 'VARCHAR(255)', '是', 'NULL', '—', '附件链接'),
        ('review_status', "ENUM('待审核','已通过','已驳回')", '否', '待审核', 'DEFAULT', '审核状态'),
    ]
)


# ===================== 5.2 约束与索引设计 =====================
doc.add_heading('5.2  约束与索引设计', level=2)

body(
    '完整性约束是数据库物理结构设计的核心组成部分，它在数据库引擎层面保障数据的正确性、'
    '一致性和完整性。本节系统梳理本数据库采用的全部约束类型及索引配置。'
)

doc.add_heading('5.2.1  PRIMARY KEY 主键约束', level=3)
body(
    '每张表均定义了PRIMARY KEY约束（见表5-1至表5-12）。主键采用两种策略：对于编码型标识符'
    '（如college_id、researcher_id、project_id），使用CHAR类型并显式赋值；对于流水型记录'
    '（如budget_id、exp_id、notice_id等），使用INT AUTO_INCREMENT自增，确保唯一性和插入性能。'
    '复合主键仅出现于review表（project_id + expert_id），对应m:n联系的关联关系模式。'
)

doc.add_heading('5.2.2  FOREIGN KEY 外键约束', level=3)
body(
    '系统中所有实体间的关联均通过FOREIGN KEY约束在数据库层面明确定义，确保引用完整性。'
    '表5-13汇总了全部外键约束及其级联策略。级联策略的选择遵循以下原则：'
    '（1）依赖型实体（如预算科目归属于项目）采用ON DELETE CASCADE — 项目删除时其预算科目'
    '自动删除，避免孤数据；'
    '（2）参照型关联（如研究人员引用学院、项目引用研究人员）采用ON DELETE SET NULL — '
    '被引用记录删除时外键字段置NULL，保留数据的可追溯性；'
    '（3）流水型关联（如经费流水引用项目）不设级联 — 历史流水记录应永久保留。'
)

make_table(
    ['序号', '子表', '外键字段', '父表', '引用字段', '级联删除', '级联更新'],
    [
        ['FK-01', 'researcher', 'college_id', 'college', 'college_id', 'SET NULL', '—'],
        ['FK-02', 'project', 'leader_id', 'researcher', 'researcher_id', '—', '—'],
        ['FK-03', 'budget', 'project_id', 'project', 'project_id', 'CASCADE', '—'],
        ['FK-04', 'expenditure', 'project_id', 'project', 'project_id', '—', '—'],
        ['FK-05', 'expenditure', 'budget_id', 'budget', 'budget_id', '—', '—'],
        ['FK-06', 'expenditure', 'operator_id', 'researcher', 'researcher_id', 'SET NULL', '—'],
        ['FK-07', 'review', 'project_id', 'project', 'project_id', 'CASCADE', '—'],
        ['FK-08', 'review', 'expert_id', 'researcher', 'researcher_id', '—', '—'],
        ['FK-09', 'notice', 'publisher_id', 'researcher', 'researcher_id', 'SET NULL', '—'],
        ['FK-10', 'contract', 'project_id', 'project', 'project_id', 'CASCADE', '—'],
        ['FK-11', 'progress_report', 'project_id', 'project', 'project_id', 'CASCADE', '—'],
        ['FK-12', 'change_request', 'project_id', 'project', 'project_id', 'CASCADE', '—'],
        ['FK-13', 'acceptance', 'project_id', 'project', 'project_id', 'CASCADE', '—'],
        ['FK-14', 'achievement', 'project_id', 'project', 'project_id', 'CASCADE', '—'],
    ]
)

doc.add_heading('5.2.3  CHECK 约束', level=3)
body(
    'CHECK约束用于在数据库层面强制执行字段取值的业务规则，防止非法数据写入。本系统中'
    '主要的CHECK约束包括：review表的score字段 CHECK(score >= 0 AND score <= 100) 限定'
    '打分范围；budget表应确保amount > 0（预算金额必须为正数）；expenditure表应确保amount > 0'
    '（交易金额必须为正数）。MySQL 8.0完整支持CHECK约束的强制执行（MySQL 5.7及之前版本'
    '仅做语法兼容而不实际执行验证），本系统基于MySQL 8.0构建，CHECK约束可正常生效。'
)

doc.add_heading('5.2.4  DEFAULT 默认值约束', level=3)
body(
    'DEFAULT约束为字段提供合理的初始值，减少数据插入时的遗漏风险。各表的主要默认值配置如下：'
    '（1）researcher.password 默认\'123456\' — 新用户创建时使用统一初始密码；'
    '（2）researcher.role 默认\'科研人员\' — 大多数新注册用户为科研人员角色；'
    '（3）project.status 默认\'申报中\' — 项目创建的初始状态；'
    '（4）budget.spent 默认0.00 — 新预算科目的已支出从零开始；'
    '（5）expenditure.approval_status 默认\'待审批\' — 经费申请提交后的初始审批状态；'
    '（6）change_request.approval_status 默认\'待审批\' — 变更申请的初始状态；'
    '（7）progress_report.review_status 默认\'待审核\' — 报告提交后的初始审核状态；'
    '（8）achievement.review_status 默认\'待审核\' — 成果登记的初始审核状态。'
)

doc.add_heading('5.2.5  UNIQUE 唯一约束', level=3)
body(
    'UNIQUE约束确保特定字段或字段组合在表中不重复。本系统的主要唯一约束包括：'
    '（1）contract.project_id 设置UNIQUE约束，保证合同与项目的一对一对应关系，'
    '防止一个项目被关联多个合同；'
    '（2）review表的复合主键(project_id, expert_id)同时也实现了UNIQUE约束的效果——'
    '每位专家对一个项目只能创建一条评审记录；'
    '（3）如未来需要，可考虑对researcher.email和researcher.phone添加UNIQUE约束，'
    '当前设计未包含此限制。'
)

doc.add_heading('5.2.6  索引配置', level=3)
body(
    '根据第4章4.4节的索引设计方案，为以下9个高频查询字段建立普通索引（BTREE）：'
    'project.leader_id、project.status、project.end_date、expenditure.project_id、'
    'expenditure.approval_status、review.expert_id、achievement.project_id、'
    'researcher.college_id、budget.project_id。索引采用InnoDB存储引擎默认的BTREE结构，'
    '适用于等值查询、范围查询和排序操作。主键索引（聚簇索引）和所有外键索引由数据库'
    '系统自动创建，不在此重复统计。'
)


# ===================== 5.3 存储引擎与字符集 =====================
doc.add_heading('5.3  存储引擎与字符集设计', level=2)

body(
    '存储引擎和字符集的选择直接影响数据库的事务支持能力、并发性能、数据完整性和国际化'
    '兼容性。本系统基于以下分析，选定InnoDB作为存储引擎、utf8mb4作为字符集。'
)

doc.add_heading('5.3.1  存储引擎选择：InnoDB', level=3)
body(
    'MySQL支持多种存储引擎，其中InnoDB和MyISAM是两种最常用的选择。本系统所有表统一选用'
    'InnoDB引擎，理由如下：（1）事务支持（ACID）— InnoDB支持COMMIT/ROLLBACK事务机制，'
    '经费审批流程（支出审批+更新预算）必须在同一事务中原子执行，触发器trg_expenditure_'
    'after_approve的逻辑也依赖事务回滚机制（SIGNAL SQLSTATE触发隐式回滚），MyISAM不'
    '支持事务，无法满足此需求；（2）行级锁 — InnoDB采用行级锁定（Row-Level Locking），'
    '在高并发场景下（50人同时操作），行锁可有效减少锁冲突、提升并发吞吐量，而MyISAM仅'
    '支持表级锁；（3）外键支持 — InnoDB是MySQL唯一支持FOREIGN KEY约束的存储引擎，本系统'
    '14个外键约束必须在InnoDB下才能生效；（4）崩溃恢复 — InnoDB通过重做日志（Redo Log）'
    '和撤销日志（Undo Log）实现自动崩溃恢复，保障数据安全。'
)

doc.add_heading('5.3.2  字符集选择：utf8mb4', level=3)
body(
    '字符集选择utf8mb4（UTF-8的完整四字节实现），排序规则选择utf8mb4_unicode_ci（大小写'
    '不敏感、基于Unicode标准的排序规则）。选择理由如下：（1）utf8mb4是MySQL中UTF-8的'
    '完整实现，支持所有Unicode字符（包括Emoji和罕见汉字），而MySQL的"utf8"字符集仅支持'
    '最多三字节的UTF-8，无法存储四字节字符；（2）unicode_ci排序规则提供不区分大小写的字符串'
    '比较，同时遵循Unicode Collation Algorithm（UCA），对于中文项目的名称检索更加友好；'
    '（3）统一字符集避免了跨表JOIN查询时的字符集转换开销和潜在乱码问题。'
)

doc.add_heading('5.3.3  数据表物理存储汇总', level=3)
body(
    '表5-14汇总了全部12张数据表的核心物理存储参数，包括存储引擎、字符集、预计初始记录数'
    '和单行估计大小，为数据库容量规划和性能调优提供依据。'
)

make_table(
    ['表名', '存储引擎', '字符集', '预计初始记录数', '单行估计大小', '索引数量(含自动)', '说明'],
    [
        ['college', 'InnoDB', 'utf8mb4', '6', '~120 B', '1', '学院维度表，数据量极小'],
        ['researcher', 'InnoDB', 'utf8mb4', '23', '~350 B', '2', '用户表，增长缓慢'],
        ['project', 'InnoDB', 'utf8mb4', '30+', '~500 B', '4', '核心表，高频读写，持续增长'],
        ['budget', 'InnoDB', 'utf8mb4', '43+', '~150 B', '2', '随项目数增长，中等增速'],
        ['expenditure', 'InnoDB', 'utf8mb4', '14+', '~300 B', '3', '高频插入，大量增长'],
        ['review', 'InnoDB', 'utf8mb4', '8+', '~200 B', '3', '每个项目多条，中等增速'],
        ['notice', 'InnoDB', 'utf8mb4', '6+', '~500 B', '1', '少量增长'],
        ['contract', 'InnoDB', 'utf8mb4', '以项目数为准', '~300 B', '2', '与项目1:1对应'],
        ['progress_report', 'InnoDB', 'utf8mb4', '每年递増', '~500 B', '2', '每年每个项目至少1条'],
        ['change_request', 'InnoDB', 'utf8mb4', '按需增长', '~400 B', '2', '按需增长'],
        ['acceptance', 'InnoDB', 'utf8mb4', '以项目数为准', '~300 B', '2', '与项目1:1对应'],
        ['achievement', 'InnoDB', 'utf8mb4', '30+', '~400 B', '2', '每项目可有多条，持续增长'],
    ]
)


# ---- Save ----
output_path = 'Paper/第5章_物理结构设计_v1.docx'
doc.save(output_path)
print(f"\n{'='*60}")
print("[DONE] Chapter 5 v1 generated successfully")
print(f"File: {output_path}")
print(f"{'='*60}")
