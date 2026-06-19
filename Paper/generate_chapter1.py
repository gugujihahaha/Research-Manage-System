# -*- coding: utf-8 -*-
"""生成第1章 引言 DOCX — 完美精修版"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# ---- Page setup ----
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ---- Style setup ----
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Cm(0.74)

h1_style = doc.styles['Heading 1']
h1_style.font.name = '黑体'
h1_style.font.size = Pt(16)
h1_style.font.bold = True
h1_style.font.color.rgb = RGBColor(0, 0, 0)
h1_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
h1_style.paragraph_format.space_before = Pt(12)
h1_style.paragraph_format.space_after = Pt(6)

h2_style = doc.styles['Heading 2']
h2_style.font.name = '黑体'
h2_style.font.size = Pt(14)
h2_style.font.bold = True
h2_style.font.color.rgb = RGBColor(0, 0, 0)
h2_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
h2_style.paragraph_format.space_before = Pt(8)
h2_style.paragraph_format.space_after = Pt(4)


def add_body_paragraph(text):
    """Add a body paragraph with correct formatting."""
    p = doc.add_paragraph(style='Normal')
    p.clear()
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_image_with_caption(image_path, caption_text, width_inches=5.5):
    """Add centered image with caption."""
    if not os.path.exists(image_path):
        print(f"  [SKIP] Image not found: {image_path}")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(f"（{caption_text}——待插入）")
        run.font.name = '宋体'
        run.font.size = Pt(10)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        return

    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.first_line_indent = Cm(0)
    p_img.paragraph_format.space_before = Pt(6)
    run_img = p_img.add_run()
    run_img.add_picture(image_path, width=Inches(width_inches))

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.first_line_indent = Cm(0)
    p_cap.paragraph_format.space_before = Pt(2)
    p_cap.paragraph_format.space_after = Pt(10)
    run_cap = p_cap.add_run(caption_text)
    run_cap.font.name = '宋体'
    run_cap.font.size = Pt(10)
    run_cap.font.bold = False
    run_cap.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    print(f"  [OK] Image inserted: {caption_text}")


# ================================================================
#  CHAPTER 1  —  CONTENTS
# ================================================================

doc.add_heading('第1章  引  言', level=1)

# -------------------- 1.1 设计背景 --------------------
doc.add_heading('1.1  设计背景', level=2)

add_body_paragraph(
    '随着国家创新驱动发展战略的纵深推进，高校作为基础研究和应用基础研究的主力军，'
    '承担的科研项目数量与经费规模持续攀升。以某省属重点高校为例，当前在研纵向项目'
    '（国家级、省部级、市厅级）逾五百项，横向项目（企事业单位委托及合作研发）逾三百项，'
    '年度科研经费总额约两亿元人民币。这些项目横跨人工智能、新能源材料、数字经济、'
    '智慧医疗等前沿领域，研究周期一至五年不等，涉及数十个学院及众多跨学科团队。'
    '如此庞大的科研体量，对该校科研管理的信息化水平提出了严峻挑战。'
)

add_body_paragraph(
    '然而，当前多数高校的科研管理仍停留在手工填报与电子表格混用的半手工模式。'
    '科研人员申报项目时，需反复填写包含大量重复信息的纸质表格并奔波于多部门签章；'
    '科研处管理人员需手工汇总申报材料、分派评审专家、统计评审结果，每轮申报周期'
    '耗费巨大人力成本；财务部门无法实时掌握项目预算的执行进度，经费超支往往在事后'
    '核算时才被发现；评审专家仅通过邮件或会议获取材料，沟通效率低下。信息孤岛导致'
    '数据在部门间难以高效流转，项目进度追踪困难，管理决策缺乏数据支撑。'
)

add_body_paragraph(
    '在此背景下，亟需开发一套覆盖科研项目全生命周期的信息化管理平台——从项目申报、'
    '专家评审、立项审批，到执行阶段的进展追踪与经费监控，再到验收结题与成果登记，'
    '实现科研管理的规范化、流程化与数据化。本课程设计选题源自《数据库原理实验》'
    '参考题目第七题——"高校科研项目管理系统"，要求综合运用数据库设计方法论、需求分析、'
    '概念结构设计（E-R图）、逻辑结构设计（关系模式与范式优化）、物理结构设计'
    '（表结构及约束）等核心知识，完成数据库系统的完整设计与前端应用开发。'
)

# Figure 1-1
add_image_with_caption(
    'Paper/Picture/需求分析示意图.png',
    '图1-1  高校科研管理现状痛点与信息化需求分析示意图'
)

# -------------------- 1.2 设计目标与意义 --------------------
doc.add_heading('1.2  设计目标与意义', level=2)

add_body_paragraph(
    '本课程设计的核心目标是完成一次"数据库建模 + 应用开发"全流程的完整工程实践。'
    '在数据库层面，严格遵循结构化设计方法论：通过详尽的系统需求分析，绘制分层数据流图'
    '（DFD）并编制数据字典（DD），精准刻画数据的来源、去向与处理逻辑；识别全部核心实体'
    '及其属性和联系，绘制局部E-R图，消解冲突后合并为全局E-R图；依据转换规则将E-R图映射'
    '为关系模式集合，逐一进行函数依赖分析与范式优化，确保所有关系模式达到第三范式（3NF），'
    '消除更新异常与数据冗余；在此基础上，设计视图以优化高频查询场景，创建索引以提升检索'
    '性能，编写存储过程与自定义函数以封装核心业务逻辑（如学院年度科研统计、预算执行预警计算），'
    '设计触发器以保障数据的完整性与一致性（如支出审批后自动更新已用预算、验收通过后自动记录日期）。'
)

add_body_paragraph(
    '在应用层面，选用 Python Flask + Vue 3 + Element Plus + MySQL 技术栈，构建一套'
    '功能完备的全栈 Web 系统。系统覆盖科研人员、科研处管理员、评审专家、财务处管理员'
    '四大用户角色，实现从项目申报、形式审查、专家评审、立项审批，到执行监控、验收结题、'
    '成果登记的全业务流程闭环。通过分层架构设计（Flask Blueprint 路由层 + Service 业务'
    '编排层 + Model 数据访问层）和 RESTful API 接口规范，保证代码的可维护性与可扩展性。'
    '最后，构造足量测试数据（每表不少于30条记录），通过系统化的功能测试与性能测试，'
    '验证系统的正确性、稳定性与响应效率。'
)

add_body_paragraph(
    '本课程设计的意义体现在三个维度。从学习维度而言，它将数据库原理课程中分散的知识点——'
    'ER建模、关系代数、SQL编程、范式理论、索引优化、存储过程、触发器——串联为有机的'
    '实践整体，使抽象理论在真实问题场景中落地生根，深化对数据库系统设计底层逻辑的理解。'
    '从技术维度而言，通过前端分离式Web应用架构（Flask后端 + Vue 3前端 + RESTful API）'
    '的集成实践，获得现代企业级信息系统的全栈开发经验，理解前后端协作机制与工程化开发流程。'
    '从应用维度而言，本系统具备部署至高校实际环境的潜力，可有效优化科研管理流程、降低行政'
    '沟通成本、增强经费使用的透明度与可控性，为管理决策提供数据化统计支撑，具有明确的现实应用价值。'
)

# Figure 1-2 — Generate system architecture overview
try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    # Configure Chinese font support
    plt.rcParams['font.sans-serif'] = ['SimHei']
    plt.rcParams['font.family'] = 'sans-serif'
    plt.rcParams['axes.unicode_minus'] = False

    fig, ax = plt.subplots(figsize=(10, 3.8))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 3.8)
    ax.axis('off')
    fig.patch.set_facecolor('#FAFBFC')
    ax.set_facecolor('#FAFBFC')

    # Title
    ax.text(5, 3.55, '高校科研项目管理系统 —— 总体设计目标架构', ha='center', va='center',
            fontsize=13, fontweight='bold', fontfamily='sans-serif', color='#1a1a2e')

    # Three architecture layers (bottom to top)
    layers_data = [
        (0.3, 0.3, 9.4, 0.75, '#E3F2FD', '数据库层  MySQL 8.0  (InnoDB + utf8mb4)',
         '核心表 11 张 / 视图 2 个 / 存储过程 2 个 / 触发器 2 个 / 自定义函数 1 个 / 索引 5 个'),
        (0.3, 1.2, 9.4, 0.75, '#E8F5E9', '业务逻辑层  Python Flask  (Blueprint 分层架构)',
         '路由层 12 个 Blueprint / 业务编排层 4 个 Service / 数据访问层 11 个 Model'),
        (0.3, 2.1, 9.4, 0.75, '#FFF8E1', '应用展示层  Vue 3 + Element Plus + ECharts + Axios',
         '工作台 / 项目申报 / 执行管理 / 验收管理 / 经费管理 / 成果管理 / 统计报表 / 管理后台'),
    ]

    for x, y, w, h, color, title_text, desc_text in layers_data:
        rect = mpatches.FancyBboxPatch(
            (x, y), w, h, boxstyle="round,pad=0.05",
            facecolor=color, edgecolor='#888888', linewidth=1.0
        )
        ax.add_patch(rect)
        ax.text(x + 0.2, y + h / 2 + 0.12, title_text, va='center', fontsize=9.5,
                fontweight='bold', fontfamily='sans-serif', color='#2c3e50')
        ax.text(x + 0.2, y + h / 2 - 0.18, desc_text, va='center', fontsize=8,
                fontfamily='sans-serif', color='#555555')

    # User roles at top
    roles = ['科研人员', '科研处\n管理员', '评审专家', '财务处\n管理员']
    role_colors = ['#1565C0', '#2E7D32', '#E65100', '#C62828']
    for i, (role, color) in enumerate(zip(roles, role_colors)):
        x_pos = 0.8 + i * 2.3
        rect = mpatches.FancyBboxPatch(
            (x_pos, 3.0), 1.9, 0.42, boxstyle="round,pad=0.04",
            facecolor=color, edgecolor='white', linewidth=1.5
        )
        ax.add_patch(rect)
        ax.text(x_pos + 0.95, 3.21, role, ha='center', va='center', fontsize=9,
                fontweight='bold', color='white', fontfamily='sans-serif')

    # Arrows from roles to layers
    for i in range(4):
        x_pos = 0.8 + i * 2.3
        ax.annotate('', xy=(x_pos + 0.95, 2.95), xytext=(x_pos + 0.95, 3.0),
                    arrowprops=dict(arrowstyle='->', color='#999999', lw=1.2))

    plt.tight_layout(pad=0.5)
    arch_img_path = 'Paper/Picture/系统设计目标架构图.png'
    plt.savefig(arch_img_path, dpi=200, bbox_inches='tight', facecolor='#FAFBFC')
    plt.close()
    print(f"  [OK] Architecture figure generated: {arch_img_path}")

    add_image_with_caption(arch_img_path, '图1-2  系统总体设计目标架构示意图', width_inches=5.5)

except Exception as e:
    print(f"  [WARN] Could not generate architecture figure: {e}")
    add_body_paragraph('（图1-2：系统总体设计目标架构示意图——待补充）')

# -------------------- 1.3 论文组织结构 --------------------
doc.add_heading('1.3  论文组织结构', level=2)

add_body_paragraph(
    '本文严格遵循数据库系统设计的标准流程组织全文内容，共分为九章，各章的内容安排如下：'
)

chapters = [
    ('第一章  引言',
     '阐述系统开发的项目背景与现实需求，明确课程设计的目标定位与多维度意义，'
     '并概述全文的章节组织与内容安排。'),
    ('第二章  需求分析',
     '详细分析四类目标用户角色及其核心业务需求，按功能模块逐项分解系统功能；'
     '绘制分层数据流图（顶层DFD、一层DFD、二层DFD），编制完整的数据字典（DD），'
     '并从性能、安全、可用性、可维护性等维度定义非功能性需求。'),
    ('第三章  数据库概念结构设计',
     '以项目申报、经费管理、成果管理等核心业务域为单位，分别绘制局部E-R图，'
     '识别实体、属性及联系类型（1:1 / 1:n / m:n）；通过消解属性冲突、命名冲突'
     '和结构冲突，将全部局部E-R图合并为全局E-R图。'),
    ('第四章  数据库逻辑结构设计',
     '遵循E-R图向关系模型的转换规则，将全局概念模型映射为关系模式集合；对每个'
     '关系模式进行函数依赖分析和范式判定，分解至第三范式（3NF）；设计数据视图以'
     '支持常用查询场景；建立索引以优化检索路径；定义存储过程与自定义函数；'
     '设计触发器以保障数据一致性。'),
    ('第五章  数据库物理结构设计',
     '确定每张数据表的物理存储结构，包括字段名、数据类型、长度、是否可空、默认值'
     '及各类约束（PRIMARY KEY / FOREIGN KEY / CHECK / UNIQUE / DEFAULT）；'
     '说明存储引擎选择（InnoDB）与字符集配置（utf8mb4）的依据，以及索引的物理存储方式。'),
    ('第六章  数据库实施',
     '给出完整的数据库建库建表SQL脚本（含详细注释），展示视图、索引、存储过程、触发器'
     '及自定义函数的创建语句，并附执行结果验证与关键说明。'),
    ('第七章  前端应用程序开发',
     '描绘系统的分层架构图（展示层 / 业务逻辑层 / 数据访问层 / 数据库层），说明技术'
     '选型依据与开发环境配置；绘制主要业务模块的流程图（登录认证、项目申报、经费审批、'
     '验收评审等）；通过界面截图展示系统的关键功能页面。'),
    ('第八章  系统测试与验证',
     '说明测试环境配置与测试计划；阐述模拟测试数据的构造方法（每表不少于30条记录，'
     '覆盖正常值、边界值和异常值）；列出功能测试用例表及执行结果；进行查询性能对比测试'
     '（有索引 / 无索引），验证存储过程与触发器的执行效率与正确性。'),
    ('第九章  总结与收获',
     '回顾课程设计完成的主要工作内容与阶段性成果；总结数据库设计方法论在工程实践中的'
     '应用体会；反思系统存在的不足并提出改进方向；分享学习过程中的收获与对未来学习及工作的启示。'),
]

for title, desc in chapters:
    p = doc.add_paragraph(style='Normal')
    p.clear()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    # Bold title part
    run_title = p.add_run(title)
    run_title.font.name = '宋体'
    run_title.font.size = Pt(12)
    run_title.font.bold = True
    run_title.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # Separator
    run_sep = p.add_run('：')
    run_sep.font.name = '宋体'
    run_sep.font.size = Pt(12)
    run_sep.font.bold = True
    run_sep.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # Description
    run_desc = p.add_run(desc)
    run_desc.font.name = '宋体'
    run_desc.font.size = Pt(12)
    run_desc.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ---- Save ----
output_path = 'Paper/第1章_引言_v3.docx'
doc.save(output_path)
print(f"\n{'='*60}")
print("  [DONE] Chapter 1 v3 has been generated")
print(f"  File: {output_path}")
print(f"{'='*60}")
