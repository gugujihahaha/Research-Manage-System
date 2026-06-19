# -*- coding: utf-8 -*-
"""生成第7章 前端应用程序开发 DOCX — 基于实际前端代码"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, sys

sys.stdout.reconfigure(encoding='utf-8', errors='replace')

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54); section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17); section.right_margin = Cm(3.17)

style = doc.styles['Normal']
style.font.name = '宋体'; style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Cm(0.74)

for lvl, sz in [(1, 16), (2, 14), (3, 13)]:
    s = doc.styles[f'Heading {lvl}']
    s.font.name = '黑体'; s.font.size = Pt(sz); s.font.bold = True
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if lvl == 1 else WD_ALIGN_PARAGRAPH.LEFT
    s.paragraph_format.space_before = Pt(12 if lvl == 1 else 8)
    s.paragraph_format.space_after = Pt(6 if lvl == 1 else 4)


def body(text):
    p = doc.add_paragraph(style='Normal'); p.clear()
    r = p.add_run(text)
    r.font.name = '宋体'; r.font.size = Pt(12)
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    return p


def code_block(code):
    p = doc.add_paragraph(style='Normal'); p.clear()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'F5F5F5'); shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    r = p.add_run(code)
    r.font.name = 'Consolas'; r.font.size = Pt(7.5)
    r.font.color.rgb = RGBColor(30, 30, 30)
    return p


def make_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.font.name = '宋体'; r.font.size = Pt(9.5); r.font.bold = True
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), '4472C4'); shd.set(qn('w:val'), 'clear')
        c._tc.get_or_add_tcPr().append(shd)
        r.font.color.rgb = RGBColor(255, 255, 255)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.rows[i + 1].cells[j]; c.text = ''
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val)); r.font.name = '宋体'; r.font.size = Pt(8.5)
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if i % 2 == 0:
                shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'F2F6FC'); shd.set(qn('w:val'), 'clear')
                c._tc.get_or_add_tcPr().append(shd)
    doc.add_paragraph()
    return t


def add_figure(img_path, caption, width=5.5):
    if not os.path.exists(img_path): return
    pi = doc.add_paragraph(); pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pi.paragraph_format.first_line_indent = Cm(0); pi.paragraph_format.space_before = Pt(6)
    pi.add_run().add_picture(img_path, width=Inches(width))
    pc = doc.add_paragraph(); pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pc.paragraph_format.first_line_indent = Cm(0)
    pc.paragraph_format.space_before = Pt(2); pc.paragraph_format.space_after = Pt(10)
    r = pc.add_run(caption); r.font.name = '宋体'; r.font.size = Pt(10)
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


# ================================================================
#  MATPLOTLIB
# ================================================================
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
plt.rcParams['font.sans-serif'] = ['SimHei']
plt.rcParams['font.family'] = 'sans-serif'
plt.rcParams['axes.unicode_minus'] = False

def save_fig(fig, path):
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close(fig)


# ================================================================
#  CHAPTER 7
# ================================================================
doc.add_heading('第7章  前端应用程序开发', level=1)

body(
    '本章在前六章数据库设计的基础上，阐述前端Web应用程序的设计与实现。系统采用前后端'
    '分离的SPA（Single Page Application）架构，前端基于Vue 3 Composition API + Element '
    'Plus UI组件库构建响应式用户界面，通过Axios HTTP客户端与后端Flask RESTful API进行'
    '数据交互。后端采用Flask Blueprint分层架构，将12个功能模块解耦为独立的路由组件。'
    '本章依次介绍系统架构设计、技术选型、核心功能模块的流程设计以及系统界面实现。'
)

# ===================== 7.1 系统架构设计 =====================
doc.add_heading('7.1  系统架构设计', level=2)

body(
    '系统采用经典的四层B/S架构，自底向上依次为数据库层、数据访问层、业务逻辑层和应用'
    '展示层。各层职责明确、接口清晰，层间通过标准化的协议或API进行通信。'
)

doc.add_heading('7.1.1  分层架构', level=3)

body(
    '（1）数据库层（MySQL 8.0）。采用InnoDB存储引擎和utf8mb4字符集，存储12张核心数据表、'
    '2个视图、2个存储过程、2个触发器、1个自定义函数和5个索引。数据库层通过PyMySQL驱动'
    '与数据访问层连接，所有SQL操作均使用参数化查询。'
)

body(
    '（2）数据访问层（Model层，11个模块）。位于backend/models/目录下，每个模块封装对应'
    '数据表的CRUD操作及复杂查询。Model层通过backend/utils/db.py提供的数据库连接池和'
    'execute_query()/execute_one()/execute_update()三个工具函数访问数据库，统一使用'
    'pymysql.cursors.DictCursor以字典格式返回查询结果。所有写操作在try/except块中执行，'
    '异常时自动回滚。'
)

body(
    '（3）业务逻辑层（Service + API层）。Service层（backend/services/，4个模块）负责'
    '业务流程编排，如项目状态流转的合法性校验、项目编号的自动生成、经费审批的事务协调等。'
    'API层（backend/api/，12个Blueprint路由模块）是前端与后端之间的RESTful接口层，每个'
    'Blueprint对应一个功能域，通过@app.route装饰器定义路由，统一返回JSON格式'
    '{"code": 200, "data": ..., "message": "..."}。所有Blueprint在app.py的create_app()'
    '工厂函数中注册，URL前缀统一为/api。'
)

body(
    '（4）应用展示层（Vue 3 + Element Plus）。前端为单页应用（SPA），单一HTML入口文件'
    'templates/index.html包含完整的CSS样式系统和Vue 3模板语法。JavaScript代码按功能'
    '拆分为20个模块文件：app.js（Vue主应用，负责登录/注册/角色路由/待办轮询）、api.js'
    '（Axios封装，统一错误拦截）、pages/目录下18个页面组件。第三方库通过CDN引入，包括'
    'Vue 3、Element Plus、ECharts图表库和Axios HTTP客户端。'
)

# ---- Figure 7-1: System Architecture ----
try:
    fig, ax = plt.subplots(figsize=(11, 6.5))
    ax.set_xlim(0, 11); ax.set_ylim(0, 6.5); ax.axis('off')
    ax.set_facecolor('#FAFBFC'); fig.patch.set_facecolor('#FAFBFC')
    ax.text(5.5, 6.2, '系统分层架构图', ha='center', va='center',
            fontsize=13, fontweight='bold', color='#1a1a2e')

    layers = [
        (0.5, 0.3, 10.0, 0.9, '#E3F2FD', '数据库层  MySQL 8.0',
         'college / researcher / project / budget / expenditure / review / notice\n'
         'contract / progress_report / change_request / acceptance / achievement\n'
         '2 视图 | 2 存储过程 | 2 触发器 | 1 函数 | 5 索引 | InnoDB + utf8mb4'),
        (0.5, 1.4, 10.0, 0.9, '#E8F5E9', '数据访问层  backend/models/ (11个Model)',
         'college / researcher / project / budget / expenditure / review\n'
         'notice / contract / progress_report / change_request / acceptance / achievement\n'
         'backend/utils/db.py: PyMySQL连接池 + DictCursor + autocommit=False'),
        (0.5, 2.5, 10.0, 1.1, '#FFF3E0', '业务逻辑层  Flask Blueprint (12路由) + Service (4编排)',
         'auth / projects / notices / reviews / reports / acceptance\n'
         'funding / achievements / budgets / contracts / stats / upload\n'
         'project_service / funding_service / acceptance_service / stats_service'),
        (0.5, 3.8, 10.0, 0.9, '#FCE4EC', '应用展示层  Vue 3 SPA  (templates/index.html)',
         'app.js (主应用+角色路由) | api.js (Axios封装) | pages/ (18个页面组件)\n'
         'Vue 3 Composition API | Element Plus | ECharts | CDN加载'),
        (0.5, 4.9, 10.0, 0.55, '#ECEFF1', '客户端  现代浏览器 (Chrome / Firefox / Edge)',
         '科研人员 | 科研处管理员 | 评审专家 | 财务处管理员'),
    ]
    colors = ['#E3F2FD','#E8F5E9','#FFF3E0','#FCE4EC','#ECEFF1']
    for (x, y, w, h, color, title_text, desc_text), ec in zip(layers, ['#1565C0','#2E7D32','#E65100','#C62828','#546E7A']):
        rect = mpatches.FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.05",
                                        facecolor=color, edgecolor=ec, linewidth=1.3)
        ax.add_patch(rect)
        ax.text(x + 0.2, y + h - 0.18, title_text, va='top', fontsize=9.5,
                fontweight='bold', color='#1a1a2e')
        ax.text(x + 0.2, y + 0.12, desc_text, va='bottom', fontsize=7, color='#555555')

    plt.tight_layout(pad=0.3)
    save_fig(fig, 'Paper/Picture/fig7_1_系统架构图.png')
    add_figure('Paper/Picture/fig7_1_系统架构图.png', '图7-1  系统分层架构图')
except Exception as e:
    print(f"WARN fig7-1: {e}")

# 7.1.2 前后端交互
doc.add_heading('7.1.2  前后端交互机制', level=3)

body(
    '前后端通过HTTP协议进行数据交互，遵循RESTful API设计规范。前端api.js封装了Axios实例，'
    '配置baseURL为http://127.0.0.1:5000/api，timeout为15秒，启用withCredentials以支持'
    '跨域Cookie传递。响应拦截器统一处理401未授权状态（清除sessionStorage并刷新页面）和'
    '网络错误，通过Element Plus的ElMessage组件弹出错误提示。后端Flask通过flask_cors扩展'
    '启用CORS（supports_credentials=True），允许前端跨域请求。'
)

body(
    'API统一返回格式为{"code": 200, "data": {...}, "message": "success"}，其中code=200'
    '表示成功，code=400表示客户端请求错误，code=500表示服务器内部错误。前端所有数据获取'
    '均通过api.get/post/put/delete方法发起，文件上传使用独立的uploadFile函数，以FormData'
    '格式发送multipart/form-data请求。'
)

# ---- Figure 7-2: Frontend-Backend Interaction ----
try:
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.set_xlim(0, 10); ax.set_ylim(0, 4); ax.axis('off')
    ax.set_facecolor('white'); fig.patch.set_facecolor('white')
    ax.text(5, 3.7, '前后端交互流程', ha='center', va='center',
            fontsize=12, fontweight='bold', color='#1a1a2e')

    # Frontend box
    rect = mpatches.FancyBboxPatch((0.5, 2.2), 3.8, 1.0, boxstyle="round,pad=0.06",
                                    facecolor='#FCE4EC', edgecolor='#C62828', linewidth=1.5)
    ax.add_patch(rect)
    ax.text(2.4, 2.7, 'Vue 3 前端 (SPA)\nElement Plus + Axios', ha='center', fontsize=9, fontweight='bold', color='#C62828')

    # Backend box
    rect = mpatches.FancyBboxPatch((5.7, 2.2), 3.8, 1.0, boxstyle="round,pad=0.06",
                                    facecolor='#FFF3E0', edgecolor='#E65100', linewidth=1.5)
    ax.add_patch(rect)
    ax.text(7.6, 2.7, 'Flask 后端\n12 Blueprint + 4 Service', ha='center', fontsize=9, fontweight='bold', color='#E65100')

    # Arrows
    ax.annotate('HTTP Request\n(GET/POST/PUT/DELETE)\nJSON / FormData',
                xy=(5.6, 3.0), xytext=(3.0, 3.0),
                arrowprops=dict(arrowstyle='->', color='#1565C0', lw=1.5),
                ha='center', va='center', fontsize=7.5, color='#1565C0')
    ax.annotate('HTTP Response\n{"code":200, "data":...}',
                xy=(3.0, 2.0), xytext=(5.6, 2.0),
                arrowprops=dict(arrowstyle='->', color='#2E7D32', lw=1.5),
                ha='center', va='center', fontsize=7.5, color='#2E7D32')

    # DB box
    rect = mpatches.FancyBboxPatch((3.5, 0.4), 3.0, 0.7, boxstyle="round,pad=0.05",
                                    facecolor='#E3F2FD', edgecolor='#1565C0', linewidth=1.2)
    ax.add_patch(rect)
    ax.text(5.0, 0.75, 'MySQL 8.0\nPyMySQL + DictCursor', ha='center', fontsize=8, color='#1565C0')
    ax.annotate('', xy=(5.0, 1.15), xytext=(5.0, 2.1),
                arrowprops=dict(arrowstyle='<->', color='#90A4AE', lw=1.0))

    plt.tight_layout(pad=0.3)
    save_fig(fig, 'Paper/Picture/fig7_2_前后端交互.png')
    add_figure('Paper/Picture/fig7_2_前后端交互.png', '图7-2  前后端交互流程')
except Exception as e:
    print(f"WARN fig7-2: {e}")


# ===================== 7.2 开发环境与技术选型 =====================
doc.add_heading('7.2  开发环境与技术选型', level=2)

body(
    '系统开发环境基于Windows 11平台，使用Visual Studio Code作为集成开发环境。'
    '表7-1列出了系统采用的全部技术组件及其版本和用途说明。'
)

make_table(
    ['层次', '技术/工具', '版本', '用途说明'],
    [
        ['数据库', 'MySQL', '8.0', '关系型数据存储，InnoDB引擎，utf8mb4字符集'],
        ['数据库驱动', 'PyMySQL', '1.1+', 'Python连接MySQL的纯Python驱动，支持DictCursor'],
        ['后端框架', 'Flask', '3.x', 'Python轻量级Web框架，Blueprint模块化路由'],
        ['跨域支持', 'Flask-CORS', '—', '处理前端跨域请求，支持Cookie传递'],
        ['安全上传', 'Werkzeug', '—', 'secure_filename安全文件命名，防止路径穿越'],
        ['前端框架', 'Vue 3', '3.x (CDN)', 'Composition API响应式框架，单页应用'],
        ['UI组件库', 'Element Plus', '2.x (CDN)', '企业级UI组件，表格/表单/对话框/标签页等'],
        ['图表库', 'ECharts', '5.x (CDN)', '统计报表的数据可视化图表（柱状图/饼图等）'],
        ['HTTP客户端', 'Axios', '1.x (CDN)', 'Promise风格的HTTP请求库，拦截器机制'],
        ['开发工具', 'VS Code', '—', '代码编辑与调试'],
        ['包管理', 'uv (pip)', '—', 'Python虚拟环境与依赖管理'],
        ['运行环境', 'Python', '3.8+', '后端运行环境'],
        ['浏览器', 'Chrome/Firefox/Edge', '现代版本', '目标运行环境'],
    ]
)


# ===================== 7.3 主要功能模块流程图 =====================
doc.add_heading('7.3  主要功能模块流程图', level=2)

body(
    '本节选取系统中三个核心业务流程——用户登录与角色路由、项目申报与审批流程、经费报销'
    '与审批流程——绘制详细的流程图，清晰展示前后端协同处理的完整逻辑链路。'
)

# ---- 7.3.1 Login Flow ----
doc.add_heading('7.3.1  用户登录与角色路由流程', level=3)

body(
    '用户登录流程涉及前端表单验证、Axios POST请求、后端认证、Session存储和Vue响应式'
    '角色路由五个环节。系统根据用户role字段动态加载对应的导航菜单（科研人员7个标签、'
    '科研处7个标签、专家4个标签、财务处5个标签），实现基于角色的界面差异化展示。'
    '待办数量通过定时轮询机制（setInterval）实时更新导航栏徽标数字。'
)

try:
    fig, ax = plt.subplots(figsize=(10, 6.5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 6.5); ax.axis('off')
    ax.set_facecolor('white'); fig.patch.set_facecolor('white')
    ax.text(5, 6.2, '用户登录与角色路由流程图', ha='center', va='center',
            fontsize=12, fontweight='bold', color='#1a1a2e')

    nodes = [
        (5, 5.6, '用户输入工号+密码', 'start'),
        (5, 4.8, '前端 Element Plus\n表单校验 (ref验证)', 'process'),
        (5, 4.0, 'api.post("/login")\n发送登录请求', 'process'),
        (5, 3.1, '后端 auth.py 验证\nresearcher_id+password', 'decision'),
        (2.5, 2.3, '返回 code=400\nElMessage.error 提示', 'error'),
        (7.5, 2.3, '返回 code=200\nuser存入sessionStorage', 'process'),
        (7.5, 1.3, 'Vue响应式更新 user ref\ncomputed导航菜单', 'process'),
        (7.5, 0.5, '按角色动态渲染\n侧边栏 + 工作台', 'end'),
    ]
    for x, y, label, ntype in nodes:
        if ntype == 'start' or ntype == 'end':
            fc, ec = '#E8F5E9', '#2E7D32'
        elif ntype == 'decision':
            fc, ec = '#FFF9C4', '#F57F17'
        elif ntype == 'error':
            fc, ec = '#FFEBEE', '#C62828'
        else:
            fc, ec = '#E3F2FD', '#1565C0'
        rect = mpatches.FancyBboxPatch((x-1.5, y-0.28), 3.0, 0.56,
                boxstyle="round,pad=0.04", facecolor=fc, edgecolor=ec, linewidth=1.2)
        ax.add_patch(rect)
        ax.text(x, y, label, ha='center', va='center', fontsize=7.5,
                fontweight='bold' if ntype in ('start','end','decision') else 'normal',
                color=ec if ntype == 'decision' else '#37474F')

    # Flow arrows
    for i in range(3):
        ax.annotate('', xy=(5, nodes[i+1][1]+0.3), xytext=(5, nodes[i][1]-0.3),
                    arrowprops=dict(arrowstyle='->', color='#546E7A', lw=1.2))
    ax.annotate('失败', xy=(2.5, 2.6), xytext=(5, 2.8),
                arrowprops=dict(arrowstyle='->', color='#C62828', lw=1.2),
                fontsize=7, color='#C62828')
    ax.annotate('成功', xy=(7.5, 2.6), xytext=(5.5, 2.8),
                arrowprops=dict(arrowstyle='->', color='#2E7D32', lw=1.2),
                fontsize=7, color='#2E7D32')
    ax.annotate('', xy=(7.5, 1.6), xytext=(7.5, 2.0),
                arrowprops=dict(arrowstyle='->', color='#546E7A', lw=1.2))
    ax.annotate('', xy=(7.5, 0.8), xytext=(7.5, 1.0),
                arrowprops=dict(arrowstyle='->', color='#546E7A', lw=1.2))

    # Role nav side note
    ax.text(1.2, 2.8, '角色菜单映射:\n科研人员→7项\n科研处→7项\n专家→4项\n财务处→5项',
            fontsize=6.5, color='#78909C', va='top')

    plt.tight_layout(pad=0.3)
    save_fig(fig, 'Paper/Picture/fig7_3_登录流程.png')
    add_figure('Paper/Picture/fig7_3_登录流程.png', '图7-3  用户登录与角色路由流程图')
except Exception as e:
    print(f"WARN fig7-3: {e}")

# ---- 7.3.2 Project Declaration Flow ----
doc.add_heading('7.3.2  项目申报与审批流程', level=3)

body(
    '项目申报与审批是系统最核心的业务流程，涉及科研人员、科研处管理员和评审专家三方协作。'
    '流程从科研人员查看通知并填写申报书开始，经形式审查→专家分配→专家评审→自动状态流转→'
    '立项审批，最终项目状态变更为"已立项"。后端project_service.py负责状态流转的合法性校验'
    '和项目编号自动生成，reviews.py实现全部专家评审完成的自动检测逻辑。'
)

try:
    fig, ax = plt.subplots(figsize=(11, 5.5))
    ax.set_xlim(0, 11); ax.set_ylim(0, 5.5); ax.axis('off')
    ax.set_facecolor('white'); fig.patch.set_facecolor('white')
    ax.text(5.5, 5.2, '项目申报与审批流程图', ha='center', va='center',
            fontsize=12, fontweight='bold', color='#1a1a2e')

    # Flow nodes - horizontal flow
    steps = [
        (0.6, 3.5, '科研人员\n浏览通知', '#E8F5E9', 'R'),
        (2.2, 3.5, '在线填报\n申报书+预算', '#E3F2FD', 'R'),
        (3.8, 3.5, '提交→\n状态: 申报中', '#FFF3E0', 'R'),
        (5.4, 3.5, '科研处\n形式审查', '#E8F5E9', 'A'),
        (5.4, 2.0, '退回修改\n→申报中', '#FFEBEE', 'A'),
        (7.0, 3.5, '分配专家\n1~N人', '#FFF3E0', 'A'),
        (8.6, 3.5, '专家评审\n打分+评语', '#E3F2FD', 'E'),
        (8.6, 2.0, '任一不通过\n→退回申报中', '#FFEBEE', 'E'),
        (10.0, 3.5, '全部通过\n→立项公示', '#E8F5E9', 'A'),
        (10.0, 4.6, '审批通过\n→已立项', '#C8E6C9', 'A'),
    ]
    for x, y, label, color, role in steps:
        w, h = 1.3, 0.65
        rect = mpatches.FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.04",
                facecolor=color, edgecolor='#607D8B', linewidth=1.0)
        ax.add_patch(rect)
        ax.text(x+w/2, y+h/2, label, ha='center', va='center', fontsize=6.5,
                fontweight='bold', color='#37474F')
        ax.text(x+w/2, y-0.15, role, ha='center', va='top', fontsize=6, color='#90A4AE')

    # Arrows for main flow
    for i in range(7):
        x1 = steps[i][0] + 1.3
        ax.annotate('', xy=(steps[i+1][0], 3.85), xytext=(x1, 3.85),
                    arrowprops=dict(arrowstyle='->', color='#546E7A', lw=1.2))

    # Branch arrows
    ax.annotate('不通过', xy=(6.0, 2.65), xytext=(6.0, 3.1),
                arrowprops=dict(arrowstyle='->', color='#C62828', lw=1.2),
                fontsize=6.5, color='#C62828')
    ax.annotate('有驳回', xy=(9.2, 2.65), xytext=(9.2, 3.1),
                arrowprops=dict(arrowstyle='->', color='#C62828', lw=1.2),
                fontsize=6.5, color='#C62828')
    ax.annotate('全部通过', xy=(10.6, 3.1), xytext=(10.6, 3.85),
                arrowprops=dict(arrowstyle='->', color='#2E7D32', lw=1.2),
                fontsize=6.5, color='#2E7D32')

    # Legend
    ax.text(0.3, 1.0, 'R=科研人员  A=科研处管理员  E=评审专家', fontsize=7, color='#90A4AE')

    plt.tight_layout(pad=0.3)
    save_fig(fig, 'Paper/Picture/fig7_4_申报流程.png')
    add_figure('Paper/Picture/fig7_4_申报流程.png', '图7-4  项目申报与审批流程图', width=6.0)
except Exception as e:
    print(f"WARN fig7-4: {e}")

# ---- 7.3.3 Funding Flow ----
doc.add_heading('7.3.3  经费报销与审批流程', level=3)

body(
    '经费报销流程涉及科研人员提交报销申请和财务处管理员审批两个阶段。科研人员选择项目'
    '的特定预算科目提交支出申请后，状态初始为"待审批"。财务处管理员审批时，后端调用'
    'funding_service.py进行业务校验，数据库触发器trg_expenditure_after_approve在底层'
    '执行预算余额检查和自动更新，超支则回滚事务。审批结果通过Element Plus的消息提示'
    '组件实时反馈给前端用户。'
)

try:
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 5); ax.axis('off')
    ax.set_facecolor('white'); fig.patch.set_facecolor('white')
    ax.text(5, 4.7, '经费报销与审批流程图', ha='center', va='center',
            fontsize=12, fontweight='bold', color='#1a1a2e')

    fnodes = [
        (5, 4.1, '科研人员选择项目\n+预算科目+金额+用途', '#E8F5E9'),
        (5, 3.2, 'POST /api/funding/expenditure\nstatus: 待审批', '#E3F2FD'),
        (5, 2.2, '财务处管理员\n审批操作', '#FFF3E0'),
        (2.5, 1.2, '驳回\nstatus: 已驳回\nElMessage.warning', '#FFEBEE'),
        (7.5, 2.2, '通过：触发\ntrg_expenditure_after_approve', '#E8F5E9'),
        (7.5, 1.2, '检查预算余额', '#FFF9C4'),
        (5.5, 0.3, '余额不足→SIGNAL\n事务回滚 拦截', '#FFEBEE'),
        (9.5, 0.3, '余额充足→\nbudget.spent更新', '#C8E6C9'),
    ]
    for x, y, label, color in fnodes:
        w, h = 2.8, 0.55
        if len(label) > 25:
            w = 3.2
        rect = mpatches.FancyBboxPatch((x-w/2, y-h/2), w, h, boxstyle="round,pad=0.04",
                facecolor=color, edgecolor='#607D8B', linewidth=1.0)
        ax.add_patch(rect)
        ax.text(x, y, label, ha='center', va='center', fontsize=7,
                fontweight='bold', color='#37474F')

    # Arrows
    ax.annotate('', xy=(5, 3.5), xytext=(5, 3.8),
                arrowprops=dict(arrowstyle='->', color='#546E7A', lw=1.2))
    ax.annotate('', xy=(5, 2.5), xytext=(5, 2.9),
                arrowprops=dict(arrowstyle='->', color='#546E7A', lw=1.2))
    ax.annotate('驳回', xy=(2.5, 1.5), xytext=(4.0, 1.9),
                arrowprops=dict(arrowstyle='->', color='#C62828', lw=1.2),
                fontsize=7, color='#C62828')
    ax.annotate('通过', xy=(7.5, 1.5), xytext=(6.0, 1.9),
                arrowprops=dict(arrowstyle='->', color='#2E7D32', lw=1.2),
                fontsize=7, color='#2E7D32')
    ax.annotate('', xy=(7.5, 0.75), xytext=(7.5, 1.9),
                arrowprops=dict(arrowstyle='->', color='#546E7A', lw=1.2))
    ax.annotate('超支', xy=(5.5, 0.6), xytext=(7.0, 0.9),
                arrowprops=dict(arrowstyle='->', color='#C62828', lw=1.2),
                fontsize=7, color='#C62828')
    ax.annotate('充足', xy=(9.5, 0.6), xytext=(8.0, 0.9),
                arrowprops=dict(arrowstyle='->', color='#2E7D32', lw=1.2),
                fontsize=7, color='#2E7D32')

    plt.tight_layout(pad=0.3)
    save_fig(fig, 'Paper/Picture/fig7_5_经费流程.png')
    add_figure('Paper/Picture/fig7_5_经费流程.png', '图7-5  经费报销与审批流程图')
except Exception as e:
    print(f"WARN fig7-5: {e}")


# ===================== 7.4 系统界面展示 =====================
doc.add_heading('7.4  系统界面展示', level=2)

body(
    '以下通过系统运行截图展示高校科研项目管理系统的关键功能界面。所有界面均基于Vue 3 + '
    'Element Plus组件库构建，采用深色侧边栏+浅色内容区的现代设计风格，支持动态响应式布局。'
    '每个界面截图均附有功能说明。'
)

# We can't take actual screenshots, but we can describe the 6 required screenshots
screenshots = [
    ('图7-6  系统登录界面',
     '展示系统的登录页面。采用渐变深色背景+居中白色卡片设计风格，包含系统Logo、标题'
     '"高校科研项目管理系统"、工号输入框、密码输入框、登录按钮和注册入口。下方显示'
     '各角色测试账号提示信息。支持登录/注册双模式切换。'),
    ('图7-7  科研人员工作台',
     '展示科研人员登录后的工作台首页。顶部导航栏显示面包屑路径，左侧深色侧边栏展示'
     '角色专属菜单（工作台/待办/项目申报/我的项目/项目验收/成果管理/经费管理）。'
     '主内容区展示四张统计卡片（我的项目数/在研项目/待办事项/成果数量）、通知公告'
     '列表和即将到期项目提醒。'),
    ('图7-8  项目申报界面',
     '展示科研人员的项目申报表单。包含项目名称输入框、项目类型下拉选择（纵向/横向）、'
     '级别选择、起止日期选择器、总预算输入和预算科目明细表格（支持动态添加/删除行）。'
     '底部提供文件上传组件和提交按钮。表单使用Element Plus的el-form组件，带有'
     '必填字段标识（红色星号）和输入格式校验。'),
    ('图7-9  科研处项目审核界面',
     '展示科研处管理员的项目审核页面，使用el-tabs标签页组件分为"形式审查"、"分配专家"'
     '、"评审追踪"、"立项审批"、"立项撤回"五个子页面。形式审查标签页中，左侧为待审查'
     '项目列表（el-table），右侧显示选中项目的详细信息（el-descriptions）和审查操作'
     '按钮（通过/不通过+审查意见）。'),
    ('图7-10  专家评审界面',
     '展示评审专家的项目评审页面。左侧为待评审项目列表（显示项目编号、名称、类型、'
     '负责人），右侧为评审表单，包含评审结果单选（通过/不通过）、打分滑块（0-100分）'
     '和评语文本框。已评审项目列表展示评审历史和当前状态。'),
    ('图7-11  统计报表界面',
     '展示科研处统计报表页面，使用ECharts图表库进行数据可视化。包含学院年度项目统计'
     '柱状图（申报数vs立项数对比）、项目类型分布饼图（纵向vs横向）、成果类型分布'
     '玫瑰图和年度经费趋势折线图。图表支持鼠标悬停显示详细数据和图例切换交互。'),
]

for caption, desc in screenshots:
    doc.add_heading(caption, level=3)
    body(desc)
    # Note: actual screenshot would be inserted here
    p = doc.add_paragraph(style='Normal'); p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f'[{caption}：请在实际运行系统时截取对应界面并替换此占位符]')
    r.font.name = '宋体'; r.font.size = Pt(10); r.font.italic = True
    r.font.color.rgb = RGBColor(150, 150, 150)
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

body(
    '以上六张界面截图覆盖了系统的核心功能模块：登录认证、工作台仪表盘、项目申报、'
    '项目审核（科研处）、专家评审和统计报表。其余功能界面（经费管理、成果管理、验收管理、'
    '通知管理、综合查询等）在前述章节中已通过流程图和功能描述进行了详细说明，其界面风格'
    '与上述截图保持一致，均基于Element Plus组件库构建，遵循统一的视觉设计规范。'
)


# ===================== 7.5 前端项目结构 =====================
doc.add_heading('7.5  前端项目结构', level=2)

body(
    '前端代码按功能模块进行文件组织，采用"一个页面一个JS文件"的模块化策略。'
    '表7-2列出了前端目录中所有JavaScript文件及其对应的功能说明。'
)

make_table(
    ['文件路径', '对应功能', '用户角色', '说明'],
    [
        ['static/js/app.js', 'Vue 3 主应用', '全体', '登录/注册/角色路由/待办轮询/侧边栏渲染'],
        ['static/js/api.js', 'Axios 封装', '全体', 'baseURL配置/响应拦截/统一错误处理/文件上传'],
        ['static/js/pages/dashboard.js', '工作台', '全体', '统计卡片/通知列表/到期提醒/快速入口'],
        ['static/js/pages/todo.js', '待办事项', '全体', '各角色待办汇总/批量操作入口'],
        ['static/js/pages/researcher_declare.js', '项目申报', '科研人员', '申报表单/预算编制/附件上传'],
        ['static/js/pages/my_projects.js', '我的项目', '科研人员', '项目列表/状态筛选/进度详情/编辑/撤回'],
        ['static/js/pages/my_acceptance.js', '项目验收', '科研人员', '验收申请/结题材料上传'],
        ['static/js/pages/my_funding.js', '经费报销', '科研人员', '报销申请/经费流水查询'],
        ['static/js/pages/achievement.js', '成果管理', '科研人员', '成果登记/成果列表/附件上传'],
        ['static/js/pages/notices_manage.js', '通知管理', '科研处', '发布通知/已发布通知列表'],
        ['static/js/pages/review_manage.js', '项目审核', '科研处', '5个标签页:形式审查/分配专家/追踪/审批/撤回'],
        ['static/js/pages/acceptance_manage.js', '验收管理', '科研处', '待验收列表/验收评审/证书发放'],
        ['static/js/pages/expert_review.js', '项目评审', '专家', '待审列表/评审表单/评审历史'],
        ['static/js/pages/funding_manage.js', '经费管理', '财务处', '待审批/到账登记/预算监控/结余处理'],
        ['static/js/pages/stats.js', '统计报表', '科研处', 'ECharts图表:学院统计/工作量/成果分布'],
        ['static/js/pages/admin.js', '全部项目', '科研处/专家', '全项目搜索/筛选/详情查看'],
        ['templates/index.html', 'HTML入口+CSS', '全体', '完整CSS设计系统/登录页/主布局/组件样式'],
    ]
)


# ---- Save ----
output_path = 'Paper/第7章_前端应用程序开发_v1.docx'
doc.save(output_path)
print(f"\n{'='*60}")
print("[DONE] Chapter 7 v1 generated successfully")
print(f"File: {output_path}")
print(f"{'='*60}")
