# -*- coding: utf-8 -*-
"""Generate Chapter 9 - Summary and Takeaways"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import sys

sys.stdout.reconfigure(encoding='utf-8', errors='replace')

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54); section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17); section.right_margin = Cm(3.17)

style = doc.styles["Normal"]
style.font.name = "宋体"; style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Cm(0.74)

for lvl, sz in [(1, 16), (2, 14), (3, 13)]:
    s = doc.styles[f"Heading {lvl}"]
    s.font.name = "黑体"; s.font.size = Pt(sz); s.font.bold = True
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if lvl == 1 else WD_ALIGN_PARAGRAPH.LEFT
    s.paragraph_format.space_before = Pt(12 if lvl == 1 else 8)
    s.paragraph_format.space_after = Pt(6 if lvl == 1 else 4)


def body(text):
    p = doc.add_paragraph(style="Normal"); p.clear()
    r = p.add_run(text)
    r.font.name = "宋体"; r.font.size = Pt(12)
    r.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    return p


# ================================================================
doc.add_heading("第9章  总结与收获", level=1)

body(
    "本课程设计以《数据库原理实验》课程为依托，选取\"高校科研项目管理系统\"这一具有明确"
    "现实背景和复杂业务逻辑的课题，完整经历了从需求分析、概念结构设计、逻辑结构设计、"
    "物理结构设计、数据库实施到前端应用开发及系统测试的全流程。在历时数周的设计与开发"
    "过程中，将数据库原理课程中分散的理论知识点——ER建模、关系代数、范式理论、SQL编程、"
    "索引优化、存储过程、触发器、事务管理——串联为一个有机的工程实践整体，获得了宝贵的"
    "全栈开发经验。以下从项目总结和学习收获两个维度进行回顾与反思。"
)

# ===================== 9.1 =====================
doc.add_heading("9.1  项目总结", level=2)

doc.add_heading("9.1.1  完成的主要工作", level=3)

body(
    "（1）系统需求分析。深入调研了高校科研管理的实际业务流程，识别了科研人员、科研处"
    "管理员、评审专家和财务处管理员四类核心用户角色，按八大功能模块逐项分解了系统功能需求，"
    "绘制了顶层、一层和二层三层数据流图（DFD），编制了包含24项数据项、14条数据流和12个"
    "数据存储的完整数据字典（DD），并从性能、安全、可用性、可维护性和数据完整性五个维度"
    "定义了非功能性需求。"
)

body(
    "（2）数据库概念结构设计。以自底向上策略，按业务域设计了三个局部E-R图（用户与项目"
    "管理、项目审批与评审管理、经费与成果管理），在消解属性冲突、命名冲突和结构冲突三类"
    "冲突后合并为全局E-R图，最终形成了包含11类实体和14个联系的概念模型。"
)

body(
    "（3）数据库逻辑结构设计。严格遵循E-R图向关系模型的五条转换规则，将全局概念模型映射"
    "为12个关系模式；对4个核心关系模式逐一进行了详尽的函数依赖分析和范式判定（均达到3NF，"
    "review关系模式达到BCNF）；设计了2个数据视图、9个索引、2个存储过程、1个自定义函数和"
    "2个触发器，构成了完整的数据库逻辑方案。"
)

body(
    "（4）数据库物理结构设计。为12张数据表逐一确定了字段名、数据类型与长度、是否可空、"
    "默认值配置及PRIMARY KEY/FOREIGN KEY/CHECK/UNIQUE/DEFAULT五类完整性约束；定义了14个"
    "外键的级联策略（CASCADE/SET NULL）；说明了InnoDB存储引擎和utf8mb4字符集的选型理由。"
)

body(
    "（5）数据库实施。使用MySQL 8.0完成了全部12张表、2个视图、2个存储过程、2个触发器、"
    "1个自定义函数和5个索引的创建；预置了覆盖全部业务状态的初始化测试数据（30个项目、"
    "23个用户、43条预算、14条经费流水、30条成果等），并通过7步验证方案确认了所有数据库"
    "对象的正确性。"
)

body(
    "（6）前端应用开发。基于Vue 3 Composition API + Element Plus + ECharts + Axios技术栈，"
    "构建了功能完整的单页Web应用（SPA），实现了基于角色的动态菜单路由（四角色差异化导航）、"
    "RESTful API数据交互、响应式界面布局和ECharts数据可视化。"
)

body(
    "（7）系统测试。设计了12项功能测试用例，覆盖了登录认证、项目申报、形式审查、专家评审、"
    "经费管理（含超支拦截）、验收管理（含触发器验证）和统计报表七大功能域，全部测试用例"
    "均顺利通过。执行了索引性能对比测试和存储过程/触发器效率测试，验证了系统满足性能指标。"
)

doc.add_heading("9.1.2  遇到的困难与解决方案", level=3)

body(
    "困难一：项目状态流转逻辑复杂。项目从申报中到验收通过共经历了10种状态的流转，"
    "且不同状态之间的合法转换路径受到严格约束（如执行中状态不允许直接跳回申报中）。"
    "解决方案：在project_service.py中实现了状态机模式的流转校验，每一次状态变更前均检查"
    "当前状态是否允许进入目标状态，杜绝非法状态跳转。同时在数据库端通过ENUM类型对status"
    "字段的取值集合进行了硬约束，形成应用层+数据库层的双重保护。"
)

body(
    "困难二：经费超支的数据一致性保障。支出审批涉及expenditure表的status更新和budget表"
    "的spent累加两个写操作，若仅依靠应用层代码控制，在并发场景下可能出现超支漏检问题。"
    "解决方案：将预算余额校验和spent更新逻辑下沉到数据库触发器trg_expenditure_after_approve"
    "中，利用MySQL的行级锁和事务机制，确保检查余额、判断是否超支、更新spent三步操作在"
    "同一事务中原子执行。超支时通过SIGNAL SQLSTATE 45000在数据库底层回滚，即使应用层"
    "存在并发缺陷也不会导致数据损坏。"
)

body(
    "困难三：E-R图绘制的规范性与可读性平衡。全局E-R图含有11个实体和14个联系，若将所有"
    "属性以椭圆形式全部画出，图面将极度拥挤且难以阅读。解决方案：采用实体矩形框内嵌属性"
    "列表的UML风格表示法，主键以红色下划线标注，在保证语义完整性的前提下大幅提升了ER图的"
    "可读性。所有E-R图使用专业绘图工具生成，确保了图元对齐和线型统一。"
)

body(
    "困难四：前端Vue 3 CDN模式下的模块化组织。由于未使用Vue CLI/Vite构建工具（采用CDN"
    "直接引入），不能使用ES Module的import/export语法，所有组件必须在全局作用域中共享状态。"
    "解决方案：通过在app.js中以Vue 3 Composition API的setup()函数创建全局响应式对象（user、"
    "currentTab等），利用JavaScript的模块加载顺序（app.js先于页面组件加载），确保所有页面"
    "组件通过闭包访问共享的Vue实例和Element Plus全局方法（ElMessage、ElMessageBox等），"
    "实现了无需构建工具的轻量级SPA架构。"
)

doc.add_heading("9.1.3  系统存在的不足与改进方向", level=3)

body(
    "（1）用户密码以明文存储于researcher表的password字段（当前默认值为123456），存在严重"
    "安全隐患。改进方向：引入bcrypt或Argon2单向哈希算法对密码进行加盐加密存储，登录验证时"
    "比对哈希值而非明文。"
)

body(
    "（2）当前系统缺乏完善的用户认证与授权机制——登录状态仅存储于前端的sessionStorage中，"
    "后端未使用JWT Token或Session机制进行请求级别的身份验证。任何知晓API端点的人均可直接"
    "调用接口。改进方向：集成Flask-JWT-Extended或Flask-Login实现基于Token的认证体系，"
    "在Blueprint路由层添加@jwt_required装饰器进行强制鉴权。"
)

body(
    "（3）前端采用CDN加载Vue 3等第三方库，虽然简化了开发流程，但在生产环境中存在加载速度"
    "不稳定和离线不可用的问题。改进方向：将前端迁移至Vue CLI或Vite构建工具，打包所有依赖"
    "为本地静态资源，同时开启代码分割（Code Splitting）和Tree Shaking优化。"
)

body(
    "（4）文件上传功能未限制文件类型和大小，存在恶意文件上传的潜在风险。改进方向：在前端"
    "和后端双重校验文件扩展名（白名单机制，如仅允许.pdf/.doc/.docx/.jpg/.png），限制单文件"
    "大小上限（如10MB），并对上传文件进行病毒扫描（如集成ClamAV）。"
)

body(
    "（5）数据库缺少操作日志和审计追踪机制，无法追溯关键数据的变更历史（如项目状态变更、"
    "经费审批操作等）。改进方向：增设log表记录关键操作的用户、时间、操作类型和变更内容，"
    "或通过MySQL的审计插件（Audit Plugin）在数据库层面捕获所有DML操作。"
)

body(
    "（6）系统集成测试不够充分，未进行压力测试（如使用JMeter模拟50+并发用户）验证系统在"
    "高负载下的稳定性。改进方向：编写自动化测试脚本（使用pytest进行后端API测试、使用"
    "Selenium进行前端E2E测试），使用Locust或JMeter执行并发压力测试，量化系统的最大并发"
    "承载能力。"
)


# ===================== 9.2 =====================
doc.add_heading("9.2  学习收获与展望", level=2)

doc.add_heading("9.2.1  对数据库设计全流程的深刻理解", level=3)

body(
    "通过本次课程设计，我深刻体会到了数据库设计并非简单的建几张表、写几条SQL，而是"
    "一个严谨的系统工程。从需求分析中绘制DFD和编制DD开始，到概念结构设计中用E-R图抽象"
    "信息世界，再到逻辑结构设计中将E-R图转化为符合3NF的关系模式集合，最后到物理结构设计"
    "中确定每一字段的具体类型和约束——这四个阶段层层递进、环环相扣，任何前序阶段的疏漏都会"
    "在后序阶段被放大。特别是范式理论在实践中的应用让我认识到，3NF并非抽象的理论教条，而是"
    "消除数据冗余和更新异常、保障数据一致性的切实工程手段——当面对一个包含12张表、14个外键"
    "的数据库时，为什么这个字段不能放在那张表里不再是一个理论问题，而是一个直接影响系统"
    "行为和性能的实践决策。"
)

doc.add_heading("9.2.2  对数据库完整性与一致性的认识", level=3)

body(
    "数据库的完整性保障不能仅依赖应用层代码，而应在数据库引擎层面建立防线。本次设计中，"
    "经费超支拦截（trg_expenditure_after_approve触发器中的SIGNAL机制）和验收日期自动记录"
    "（trg_acceptance_update触发器）两个场景让我深刻认识到触发器在数据一致性保障中的不可"
    "替代作用——应用层代码可能因并发缺陷、异常处理遗漏或开发者疏忽而失效，但数据库触发器"
    "作为引擎内置的强制机制，只要被正确定义，就永远在对应的DML操作发生时可靠执行。此外，"
    "FOREIGN KEY约束的级联策略选择（CASCADE vs SET NULL）看似是简单的配置项，实则反映了"
    "对业务语义的深层理解——项目删除后其预算科目是否还有存在价值？这类问题必须在设计阶段"
    "明确回答。"
)

doc.add_heading("9.2.3  全栈开发的工程经验", level=3)

body(
    "在技术层面，本次课程设计提供了一个难得的全栈开发实践机会。从Python Flask后端的"
    "Blueprint分层架构（12个路由模块 + 4个Service编排模块 + 11个Model数据访问模块），"
    "到Vue 3前端的Composition API响应式编程模型，再到MySQL数据库的存储过程、触发器和"
    "函数编程——三端的协同工作需要统一的接口契约（JSON格式的RESTful API）和清晰的职责"
    "边界。我认识到，优秀的系统架构不在于使用了多么前沿的技术，而在于每一层的职责划分"
    "是否清晰、层间的接口是否简洁、异常处理是否完备。例如，Flask的全局异常处理注册"
    "（register_error_handlers）和Axios的响应拦截器共同构建了从后端到前端的统一错误"
    "处理链路，使得系统在任何异常情况下都能向用户提供友好的反馈而非技术性报错信息。"
)

doc.add_heading("9.2.4  未来展望", level=3)

body(
    "本次课程设计为我在数据库系统设计和Web全栈开发领域打下了坚实的基础。在未来的学习和"
    "工作中，我计划从以下方向深化和发展本次设计的成果："
)

body(
    "（1）技术深化方向。学习Redis缓存技术，在Flask后端与MySQL之间加入缓存层，将高频查询"
    "（如通知公告、统计报表）的结果缓存至Redis，减少数据库压力；学习Docker容器化技术，"
    "编写Dockerfile和docker-compose.yml实现系统的一键部署和环境一致性；学习CI/CD持续集成"
    "流程，使用GitHub Actions或Jenkins自动化执行测试和部署。"
)

body(
    "（2）业务拓展方向。在当前的科研项目管理功能基础上，进一步拓展科研社交功能（如"
    "研究人员之间的项目协作邀请、成果分享与引用追踪）、科研经费的精细化预算模板功能、"
    "基于机器学习的历史项目推荐（根据研究人员的研究方向和过往项目，推荐合适的申报通知"
    "和潜在合作伙伴）。"
)

body(
    "（3）学术研究方向。本次设计中使用的规范化理论（函数依赖、范式分解）可以进一步与"
    "数据库理论的学术前沿结合——如研究基于近似函数依赖的数据清洗算法、探索NoSQL数据库"
    "（MongoDB、Neo4j图数据库）在科研人员合作网络分析中的适用性、以及区块链技术在科研"
    "成果确权和学术不端防范中的应用潜力。"
)

body(
    "总之，《数据库原理实验》课程设计不仅是一次技术实践，更是一次系统思维和工程素养的"
    "全面训练。从一张空白的E-R图纸到一个可运行的全栈Web系统，每一步都凝结着对数据库原理"
    "的理解和对软件工程的敬畏。这份经历将成为我未来在计算机科学领域继续探索的宝贵基石。"
)

# ---- Save ----
output_path = "Paper/第9章_总结与收获_v1.docx"
doc.save(output_path)
print(f"\n{'='*60}")
print("[DONE] Chapter 9 v1 generated successfully")
print(f"File: {output_path}")
print(f"{'='*60}")
